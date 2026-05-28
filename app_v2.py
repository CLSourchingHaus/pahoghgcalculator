import streamlit as st
import pandas as pd
import airportsdata
from geopy.distance import geodesic
import re
import folium
from streamlit_folium import st_folium
import os
import requests
import io
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
import tempfile
import plotly.express as px
from fuzzywuzzy import process
import plotly.graph_objects as go
import datetime
from docx import Document
from copy import deepcopy

# Initialize airports data
airports = airportsdata.load('IATA')

# ==================== CONFIGURATION ====================
# USE_GITHUB_DATA controls how reference data is loaded.
# False (default): files are read from the same directory as this script.
#   Use this for local development AND for Streamlit Cloud (commit the data
#   files to the repo alongside this .py file).
# True: files are fetched from raw GitHub URLs below.
#   Use only if data files are NOT committed to the repo.
#   Update GITHUB_*_URL to your actual repo before enabling.
USE_GITHUB_DATA = False
BASE_DIR = os.path.dirname(__file__)
LOCAL_GEO_MASTER_PATH = os.path.join(BASE_DIR, 'geo_master.xlsx')
GITHUB_GEO_MASTER_URL = "https://raw.githubusercontent.com/yourusername/yourrepo/main/data/geo_master.xlsx"
LOCAL_EPA_EF_PATH = os.path.join(BASE_DIR, 'PAHO EF_usaepa_exiobase.xlsx')
GITHUB_EPA_EF_URL = "https://raw.githubusercontent.com/yourusername/yourrepo/main/data/PAHO_EF_usaepa_exiobase.xlsx"

# ==================== AWB Functions ====================
def clean_gross_weight(value):
    """Clean and standardize gross weight values. 
    
    Accepts both numeric values (float/int) and strings:
    - Numeric input (e.g., 1500, 1500.5): Assumed to be in KG, returns (weight, 'KG')
    - String input (e.g., '9320.000 K/Q', '1500 LB'): Parses value and unit
    
    Returns:
        tuple: (weight_value_in_kg, unit_extracted) or (None, None) if invalid
    """
    if pd.isna(value):
        return None, None
    
    # Check if value is already numeric (int or float)
    if isinstance(value, (int, float)):
        try:
            weight = float(value)
            # Numeric values are assumed to be in KG
            return weight, 'KG'
        except (ValueError, TypeError):
            return None, None
    
    # Handle string input
    value_str = str(value).strip().upper()
    
    # Handle cases like '9320.000 K/Q', '352K 0', '659.000 K/Q', '1500 LB'
    # Extract numeric part and unit
    match = re.match(r'^([\d,.]+)\s*([A-Z]*)[^A-Z0-9]*$', value_str)
    if not match:
        # Try alternative patterns
        match = re.match(r'^([\d,.]+)\s*([A-Z]+)\s*[/\s]*[A-Z0-9]*$', value_str)
        if not match:
            return None, None
    
    numeric_part = match.group(1)
    unit_part = match.group(2) if len(match.groups()) > 1 else ''
    
    # Determine unit
    unit = None
    if any(u in unit_part for u in ['K', 'KG', 'KGS']):
        unit = 'KG'
    elif any(u in unit_part for u in ['LB', 'LBS', 'POUND']):
        unit = 'LB'  # Add pound unit recognition
    elif unit_part:
        unit = unit_part
        
    # Handle European decimal format
    is_european = False
    if ',' in numeric_part and '.' in numeric_part:
        if numeric_part.index('.') < numeric_part.index(','):
            is_european = True
    elif ',' in numeric_part:
        parts = numeric_part.split(',')
        if len(parts) == 2 and len(parts[1]) <= 3:
            is_european = True
            
    if is_european:
        numeric_part = numeric_part.replace('.', '').replace(',', '.')
    else:
        numeric_part = numeric_part.replace(',', '')
    
    try:
        weight = float(numeric_part)
        return weight, unit
    except ValueError:
        return None, None

def normalize_apo_number(apo):
    """Normalize APO/PO identifiers for robust matching."""
    if pd.isna(apo):
        return ''
    apo_str = str(apo).upper().strip()
    apo_str = re.sub(r'\bAPO24[-_]?\b', '', apo_str)
    apo_str = re.sub(r'\bAPO[-_]?\b', '', apo_str)
    apo_str = re.sub(r'[^A-Z0-9]', '', apo_str)
    apo_str = apo_str.lstrip('0')
    return apo_str


def calculate_distances(flights):
    """Calculate distances for flight segments"""
    results = []
    total_distance = 0
    for i, flight in enumerate(flights, 1):
        row_data = {
            'leg': i,
            'origin': flight['origin'],
            'destination': flight['destination'],
            'airline': flight.get('airline', None),
            'origin_lat': None,
            'origin_lon': None,
            'dest_lat': None,
            'dest_lon': None,
            'distance_km': None
        }
        try:
            origin_coords = get_airport_coords(flight['origin'])
            dest_coords = get_airport_coords(flight['destination'])
            if None not in origin_coords and None not in dest_coords:
                distance = geodesic(origin_coords, dest_coords).kilometers
                row_data.update({
                    'origin_lat': origin_coords[0],
                    'origin_lon': origin_coords[1],
                    'dest_lat': dest_coords[0],
                    'dest_lon': dest_coords[1],
                    'distance_km': round(distance, 2)
                })
                total_distance += distance
        except Exception as e:
            print(f"Error calculating leg {i}: {e}")
        results.append(row_data)
    return results, round(total_distance, 2)

def parse_flight_route(route_str):
    """Extract flight segments from route string"""
    flights = []
    if pd.isna(route_str):
        return flights
    segments = [s.strip() for s in re.split(r'Flight \d+:', route_str) if s.strip()]
    for segment in segments:
        try:
            origin = re.search(r'Origin - ([^,]+)', segment).group(1).strip()
            dest = re.search(r'Destination - ([^,]+)', segment).group(1).strip()
            airline = re.search(r'Airline - ([^;]+)', segment)
            airline = airline.group(1).strip() if airline else None
            flights.append({'origin': origin, 'destination': dest, 'airline': airline})
        except (AttributeError, IndexError):
            print(f"Warning: Could not parse segment: {segment}")
    return flights

def get_airport_coords(location):
    """Get coordinates for airport code or city name"""
    if not location or pd.isna(location):
        return (None, None)
    location = str(location).upper().strip()
    if location in airports:
        return (airports[location]['lat'], airports[location]['lon'])
    matches = [ap for ap in airports.values() if ap['city'].upper() == location]
    if matches:
        return (matches[0]['lat'], matches[0]['lon'])
    clean_loc = re.sub(r'\b(AIRPORT|INTL|INTERNATIONAL|APT|ARPT)\b', '', location, flags=re.IGNORECASE).strip()
    if clean_loc != location:
        return get_airport_coords(clean_loc)
    return (None, None)

def get_spend_category_and_temp_control(apo, spend_df, reefer_df):
    """Get both spend category and temperature control status for a specific APO number"""
    try:
        apo_str = normalize_apo_number(apo)

        if spend_df is None or 'PurchaseOrderNumber' not in spend_df.columns:
            return None, None

        spend_df = spend_df.copy()
        spend_df['PO_Clean'] = spend_df['PurchaseOrderNumber'].astype(str).apply(normalize_apo_number)

        exact_match = spend_df[spend_df['PO_Clean'] == apo_str]

        if not exact_match.empty:
            spend_category = exact_match['SpendCategory'].iloc[0]
        else:
            formats_to_try = [
                apo_str.replace('APO', '').replace('-', '').replace(' ', '').strip(),
                apo_str.replace('APO24-', '').replace('APO', '').strip(),
                apo_str.zfill(10),
                apo_str.lstrip('0'),
                apo_str.upper(),
                apo_str.lower()
            ]
            for fmt in formats_to_try:
                fmt_match = spend_df[spend_df['PO_Clean'] == fmt]
                if not fmt_match.empty:
                    spend_category = fmt_match['SpendCategory'].iloc[0]
                    break
            else:
                partial_match = spend_df[spend_df['PO_Clean'].str.contains(apo_str, na=False)]
                if not partial_match.empty:
                    spend_category = partial_match['SpendCategory'].iloc[0]
                else:
                    return None, None

        if pd.isna(spend_category):
            return None, None

        spend_category_clean = str(spend_category).strip().upper()

        if reefer_df is None or 'SpendCategory' not in reefer_df.columns:
            return spend_category, None

        reefer_df = reefer_df.copy()
        reefer_df['SpendCategory_Clean'] = reefer_df['SpendCategory'].astype(str).str.strip().str.upper()
        temp_control_map = reefer_df.drop_duplicates('SpendCategory_Clean').set_index('SpendCategory_Clean')['temp_control'].to_dict()
        temperature_control = temp_control_map.get(spend_category_clean, None)

        return spend_category, temperature_control

    except Exception as e:
        return None, None

def process_awb_file(file_path, sheet_name='AWB', spend_file=None, geo_content=None, use_github=False):
    """Process Excel file with AWB data and add temperature control + spend category"""
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=1)
        if 'Route/Stops' not in df.columns or 'Gross weight' not in df.columns:
            st.error("Required columns not found. Need both 'Route/Stops' and 'Gross weight'")
            return None
        
        # Load spend and geo data for temperature control if provided
        spend_df = None
        reefer_df = None
        
        if spend_file and geo_content:
            try:
                # Extract spend data
                spend_df = extract_spend_data(spend_file)
                # Extract temperature control mapping from geo_master
                geo_data = get_geo_master_data(geo_content, use_github=use_github)
                if geo_data[0] is not None:
                    _, _, _, reefer_df, _, _ = geo_data
            except Exception as e:
                st.warning(f"Could not load temperature control data: {e}")
        
        all_results = []
        for idx, row in df.iterrows():
            try:
                excel_row_num = idx + 3
                apo = row['APO'] if 'APO' in row.index else (row.iloc[0] if len(row) > 0 else None)
                supplier = row['Supplier'] if 'Supplier' in row.index else (row.iloc[2] if len(row) > 2 else None)
                
                if pd.isna(row['Route/Stops']):
                    continue
                    
                flights = parse_flight_route(row['Route/Stops'])
                if not flights:
                    continue
                
                # NEW APPROACH: Calculate distance from first origin to final destination
                first_origin = flights[0]['origin']
                final_destination = flights[-1]['destination']
                
                # Get coordinates for first origin and final destination
                origin_coords = get_airport_coords(first_origin)
                dest_coords = get_airport_coords(final_destination)
                
                total_distance = 0
                if None not in origin_coords and None not in dest_coords:
                    total_distance = geodesic(origin_coords, dest_coords).kilometers
                
                # Clean and convert gross weight
                gross_weight = row['Gross weight']
                clean_weight, uom = clean_gross_weight(gross_weight)
                
                gross_weight_ton = None
                if clean_weight is not None and uom is not None:
                    if uom.upper() in ['KG', 'KGS', 'K']:
                        gross_weight_ton = clean_weight / 1000
                    elif uom.upper() in ['LB', 'LBS', 'POUND']:
                        gross_weight_ton = clean_weight / 2204.62
                    else:
                        gross_weight_ton = clean_weight
                
                # Get spend category and temperature control if data is available
                spend_category = None
                temperature_control = None
                
                if spend_df is not None and reefer_df is not None and apo is not None:
                    spend_category, temperature_control = get_spend_category_and_temp_control(apo, spend_df, reefer_df)
                
                # Calculate emissions based on total distance (not individual legs)
                total_ghg_g = 0
                if total_distance > 0 and gross_weight_ton is not None:
                    # Use appropriate emission factor based on distance
                    emission_factor = 1363 if total_distance < 1500 else 1272  #1272 or 1363 uniform for all flights
                    total_ghg_g = emission_factor * gross_weight_ton * total_distance
                
                # Create a single result for the entire shipment (no individual legs)
                result = {
                    'awb_row': excel_row_num,
                    'APO': apo,
                    'Supplier': supplier,
                    'leg': 'DIRECT',
                    'origin': first_origin,
                    'destination': final_destination,
                    'airline': 'Multiple',  # Since we're combining multiple flights
                    'origin_lat': origin_coords[0] if origin_coords else None,
                    'origin_lon': origin_coords[1] if origin_coords else None,
                    'dest_lat': dest_coords[0] if dest_coords else None,
                    'dest_lon': dest_coords[1] if dest_coords else None,
                    'distance_km': round(total_distance, 2),
                    'gross_weight': clean_weight,
                    'UoM': uom,
                    'gross_weight_ton': gross_weight_ton,
                    'E.F.': 1363 if total_distance < 1500 else 1272 if total_distance > 0 else None,
                    'Units': 'g CO2e/t-km',
                    'Spend Category': spend_category,
                    'Temperature Control': temperature_control,
                    'ghg_emissions_gCO2e': total_ghg_g if total_ghg_g != 0 else None,
                    'ghg_emissions_tCO2e': total_ghg_g / 1000000 if total_ghg_g != 0 else None
                }
                
                all_results.append(result)
                
            except Exception as e:
                print(f"Error processing row {idx + 3}: {e}")
                continue
        
        if not all_results:
            st.warning("No valid flight data found in file")
            return None
            
        results_df = pd.DataFrame(all_results)
        column_order = [
            'APO', 'Supplier', 'awb_row', 'leg', 'origin', 'destination', 'airline',
            'gross_weight', 'UoM', 'gross_weight_ton', 'distance_km',
            'E.F.', 'Units', 'ghg_emissions_gCO2e', 'ghg_emissions_tCO2e',
            'Spend Category', 'Temperature Control', 'origin_lat', 'origin_lon', 'dest_lat', 'dest_lon'
        ]
        column_order = [col for col in column_order if col in results_df.columns]
        
        return results_df[column_order]
        
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def show_combined_map(awb_data, bol_data, precomputed_mappings):
    """Show combined air and ocean routes on map"""
    if (awb_data is None or awb_data.empty) and (bol_data is None or bol_data.empty):
        st.warning("No transport data available for mapping")
        return
    
    # Create base map
    avg_lat, avg_lon = 5, 0  # Default center if no data
    m = folium.Map(
        location=[avg_lat, avg_lon],
        zoom_start=2.4,
        tiles='CartoDB positron',
        control_scale=True
    )
    
    # Process AWB data (air routes - orange)
    if awb_data is not None and not awb_data.empty:
        # Ensure required numeric columns exist to avoid KeyError during aggregation
        awb_data = awb_data.copy()
        if 'ghg_emissions_tCO2e' not in awb_data.columns:
            if 'ghg_emissions_gCO2e' in awb_data.columns:
                awb_data['ghg_emissions_tCO2e'] = pd.to_numeric(awb_data['ghg_emissions_gCO2e'], errors='coerce').fillna(0) / 1_000_000
            else:
                awb_data['ghg_emissions_tCO2e'] = 0.0
        # Ensure distance_km is numeric
        awb_data['distance_km'] = pd.to_numeric(awb_data.get('distance_km', 0), errors='coerce').fillna(0)
        awb_routes = awb_data.groupby([
            'origin', 'origin_lat', 'origin_lon',
            'destination', 'dest_lat', 'dest_lon'
        ]).agg({
            'APO': lambda x: ', '.join(sorted(set(x.astype(str)))) if len(x) > 0 else 'N/A',
            'Supplier': lambda x: ', '.join(sorted(set(x.astype(str)))) if len(x) > 0 else 'N/A',
            'distance_km': 'first',
            'ghg_emissions_tCO2e': 'sum',
            'leg': 'count'
        }).reset_index().rename(columns={'leg': 'count'})
        
        for _, route in awb_routes.iterrows():
            popup_content = f"""
            <div style="font-family: Arial; font-size: 14px; width: 250px">
                <h4 style="margin-bottom: 5px; color: #e67e22;">Air Route</h4>
                <hr style="margin: 5px 0;">
                <p style="margin: 3px 0;"><b>Route:</b> {route['origin']} → {route['destination']}</p>
                <p style="margin: 3px 0;"><b>Flights:</b> {route['count']}</p>
                <p style="margin: 3px 0;"><b>Distance:</b> {route['distance_km']:,.0f} km</p>
                <p style="margin: 3px 0;"><b>Total Emissions:</b> {route['ghg_emissions_tCO2e']:,.2f} tCO2e</p>
                <p style="margin: 3px 0;"><b>APOs:</b> {route['APO']}</p>
                <p style="margin: 3px 0;"><b>Suppliers:</b> {route['Supplier']}</p>
            </div>
            """
            
            folium.PolyLine(
                locations=[
                    [route['origin_lat'], route['origin_lon']],
                    [route['dest_lat'], route['dest_lon']]
                ],
                color='#e67e22',  # Orange for air
                weight=2,  # Reduced from 3 to make lines thinner
                opacity=0.2,  # Reduced from 0.7 to make more transparent
                popup=folium.Popup(popup_content, max_width=300),
                tooltip=f"Air: {route['origin']} → {route['destination']}"
            ).add_to(m)
    
    # Process BOL data (ocean routes - blue)
    if bol_data is not None and not bol_data.empty:
        # Use precomputed city_to_coords mapping
        city_to_coords = precomputed_mappings['city_to_coords']
        # Ensure required numeric columns exist to avoid KeyError during aggregation
        bol_data = bol_data.copy()
        if 'ghg_emissions_tCO2e' not in bol_data.columns:
            if 'ghg_emissions_gCO2e' in bol_data.columns:
                bol_data['ghg_emissions_tCO2e'] = pd.to_numeric(bol_data['ghg_emissions_gCO2e'], errors='coerce').fillna(0) / 1_000_000
            else:
                bol_data['ghg_emissions_tCO2e'] = 0.0
        # Ensure Sea Distance (km) is numeric
        bol_data['Sea Distance (km)'] = pd.to_numeric(bol_data.get('Sea Distance (km)', 0), errors='coerce').fillna(0)
        
        bol_routes = bol_data.groupby([
            'Port of loading', 'Port of discharge'
        ]).agg({
            'APO no.': lambda x: ', '.join(sorted(set(x.astype(str)))) if len(x) > 0 else 'N/A',
            'Shipper name': lambda x: ', '.join(sorted(set(x.astype(str)))) if len(x) > 0 else 'N/A',
            'Sea Distance (km)': 'first',
            'ghg_emissions_tCO2e': 'sum',
            'Measurement': 'count'
        }).reset_index().rename(columns={'Measurement': 'count'})
        
        for _, route in bol_routes.iterrows():
            origin = route['Port of loading'].upper()
            dest = route['Port of discharge'].upper()
            
            origin_coords = city_to_coords.get(origin)
            dest_coords = city_to_coords.get(dest)
            
            if origin_coords and dest_coords:
                popup_content = f"""
                <div style="font-family: Arial; font-size: 14px; width: 250px">
                    <h4 style="margin-bottom: 5px; color: #3498db;">Ocean Route</h4>
                    <hr style="margin: 5px 0;">
                    <p style="margin: 3px 0;"><b>Route:</b> {origin} → {dest}</p>
                    <p style="margin: 3px 0;"><b>Shipments:</b> {route['count']}</p>
                    <p style="margin: 3px 0;"><b>Distance:</b> {route['Sea Distance (km)']:,.0f} km</p>
                    <p style="margin: 3px 0;"><b>Total Emissions:</b> {route['ghg_emissions_tCO2e']:,.2f} tCO2e</p>
                    <p style="margin: 3px 0;"><b>APOs:</b> {route['APO no.']}</p>
                    <p style="margin: 3px 0;"><b>Suppliers:</b> {route['Shipper name']}</p>
                </div>
                """
                
                folium.PolyLine(
                    locations=[
                        [origin_coords['lat'], origin_coords['lng']],
                        [dest_coords['lat'], dest_coords['lng']]
                    ],
                    color='#3498db',  # Blue for ocean
                    weight=2,
                    opacity=0.6,
                    popup=folium.Popup(popup_content, max_width=300),
                    tooltip=f"Ocean: {origin} → {dest}"
                ).add_to(m)
    
    # Add legend
    legend_html = """
    <div style="position: fixed; bottom: 50px; left: 50px; width: 180px;
                z-index: 1000; background-color: white; padding: 10px;
                border: 2px solid grey; border-radius: 5px; font-size: 14px;
                font-family: Arial; box-shadow: 3px 3px 5px rgba(0,0,0,0.2)">
        <h4 style="margin: 0 0 8px 0; padding: 0;">Route Types</h4>
        <div style="display: flex; align-items: center; margin-bottom: 5px;">
            <div style="background: #e67e22; height: 20px; width: 20px; 
                        margin-right: 10px; opacity: 0.5;"></div>
            <span>Air Transport</span>
        </div>
        <div style="display: flex; align-items: center;">
            <div style="background: #3498db; height: 20px; width: 20px; 
                        margin-right: 10px; opacity: 0.4;"></div>
            <span>Ocean Transport</span>
        </div>
    </div>
    """
    m.get_root().html.add_child(folium.Element(legend_html))
    
    st_folium(m, width=1200, height=700, returned_objects=[])

def style_awb_dataframe(df):
    """Apply styling to highlight empty cells in flight legs"""
    def highlight_empty(val):
        if pd.isna(val) or val == '':
            return 'background-color: yellow'
        return ''
    
    # Apply styling only to flight legs (not TOTAL rows)
    styled_df = df.copy()
    if 'leg' in styled_df.columns:
        flight_legs_mask = styled_df['leg'].apply(lambda x: isinstance(x, int) or (isinstance(x, str) and x.isdigit()))
        styled_df = styled_df.style.map(highlight_empty, subset=pd.IndexSlice[flight_legs_mask, :])
    
    return styled_df

def create_excel_download(df, file_name):
    """Create an Excel file with highlighting for empty cells"""
    wb = Workbook()
    ws = wb.active
    
    # Create yellow fill for highlighting
    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
    
    # Write dataframe to worksheet
    for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
        for c_idx, value in enumerate(row, 1):
            ws.cell(row=r_idx, column=c_idx, value=value)
            
            # Highlight empty cells in flight legs (not TOTAL rows)
            if r_idx > 1:  # Skip header
                leg_value = df.iloc[r_idx-2]['leg'] if 'leg' in df.columns else None
                if (pd.isna(value) or value == '') and (isinstance(leg_value, int) or (isinstance(leg_value, str) and leg_value.isdigit())):
                    ws.cell(row=r_idx, column=c_idx).fill = yellow_fill
    
    # Save to temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        wb.save(tmp.name)
        tmp.seek(0)
        data = tmp.read()
    
    # Create download button
    st.download_button(
        label=f"Download {file_name}",
        data=data,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    
    # Clean up
    os.unlink(tmp.name)

# ==================== BOL Functions ====================
def clean_bol_weight(value):
    """Clean and standardize BOL gross weight values.
    
    Accepts both numeric values (float/int) and strings:
    - Numeric input (e.g., 1500, 1500.5): Assumed to be in KG, returns value directly
    - String input: Parses numeric value from string (units are ignored, assumed KG)
    
    Returns:
        float: weight in kg, or None if invalid
    """
    if pd.isna(value):
        return None
    
    # Check if value is already numeric (int or float)
    if isinstance(value, (int, float)):
        try:
            weight = float(value)
            # Return the weight directly, assumed to be in KG
            return weight if weight > 0 else None
        except (ValueError, TypeError):
            return None
    
    # Handle string input
    value_str = str(value).strip().upper()
    
    # Handle empty strings
    if not value_str:
        return None
    
    # Extract numeric part using regex that captures numbers with commas and decimals
    # This pattern handles: digits, commas, decimal points, and optional negative sign
    match = re.search(r'[-]?[\d,]+(?:\.\d+)?', value_str)
    if not match:
        return None
    
    numeric_part = match.group(0)
    
    # Handle European decimal format (comma as decimal separator)
    is_european = False
    if ',' in numeric_part and '.' in numeric_part:
        # If both comma and period are present, check which is the decimal separator
        comma_pos = numeric_part.find(',')
        dot_pos = numeric_part.find('.')
        if comma_pos > dot_pos:  # Format like "1.000,00" (European)
            is_european = True
    elif ',' in numeric_part and numeric_part.count(',') == 1:
        # Single comma - check if it's likely a decimal separator
        parts = numeric_part.split(',')
        if len(parts) == 2 and len(parts[1]) <= 3:
            is_european = True
            
    if is_european:
        numeric_part = numeric_part.replace('.', '').replace(',', '.')
    else:
        numeric_part = numeric_part.replace(',', '')
    
    try:
        weight = float(numeric_part)
        return weight if weight > 0 else None
    except ValueError:
        return None

def extract_bol_data(file_path, sheet_name='BoL'):
    """Extracts specific columns from the BoL worksheet"""
    target_columns = [
        'APO no.', 'Port of loading', 'Port of discharge',
        'Shipper name', 'Shipper address', 'Consignee',
        'Consignee address', 'Gross weight, kg', 'Measurement'
    ]
    
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=1)
        extracted_df = df[target_columns].copy()
        
        # Clean the Gross weight column using our new function
        extracted_df['Gross weight, kg'] = extracted_df['Gross weight, kg'].apply(clean_bol_weight)
        
        return extracted_df
    except Exception as e:
        st.error(f"Error processing BoL data: {e}")
        return None

@st.cache_data
def load_and_process_geo_master(_geo_content, use_github=False):
    """Load and preprocess geo master data with caching, supporting both GitHub and local sources"""
    try:
        if use_github:
            # _geo_content is bytes from GitHub - convert to BytesIO
            geo_file_obj = io.BytesIO(_geo_content)
        else:
            # _geo_content could be a file path (str) or BytesIO object
            if isinstance(_geo_content, str):
                # Local file path - read it into BytesIO
                with open(_geo_content, 'rb') as f:
                    geo_file_obj = io.BytesIO(f.read())
            elif isinstance(_geo_content, io.BytesIO):
                # Already BytesIO object - use as is
                geo_file_obj = _geo_content
            elif isinstance(_geo_content, bytes):
                # bytes object - convert to BytesIO
                geo_file_obj = io.BytesIO(_geo_content)
            else:
                st.error(f"Unsupported geo_content type: {type(_geo_content)}")
                return None
        
        # Now geo_file_obj is always a BytesIO object with seek capability
        geo_df = pd.read_excel(
            geo_file_obj,
            sheet_name='geo',
            usecols=['city_ascii', 'lat', 'lng', 'country', 'iso3', 'continent']
        )
        geo_file_obj.seek(0)  # Reset to beginning
        
        cerdi_df = pd.read_excel(
            geo_file_obj,
            sheet_name='cerdi',
            usecols=['iso1', 'iso2', 'seadistance_km']
        )
        geo_file_obj.seek(0)
        
        ef_df = pd.read_excel(
            geo_file_obj,
            sheet_name='ef',
            usecols=['Port of Origin', 'Port of Destination', 
                    'Corresponding trade line group', 'Reefer', 'Dry']
        )
        geo_file_obj.seek(0)
        
        reefer_df = pd.read_excel(
            geo_file_obj,
            sheet_name='temperature',
            usecols=['SpendCategory', 'temp_control']
        )
        geo_file_obj.seek(0)
        
        supplier_df = pd.read_excel(
            geo_file_obj,
            sheet_name='supplier',
            usecols=['Supplier', 'country', 'supplier_specific_ef_sc123']
        )
        
        # Preprocess data for faster access
        # 1. Geo data: create optimized mappings
        geo_unique = geo_df.drop_duplicates(subset=['city_ascii'], keep='first').copy()
        geo_unique['city_ascii_upper'] = geo_unique['city_ascii'].str.upper()
        city_to_coords = {}
        for _, row in geo_unique.iterrows():
            city = row['city_ascii']
            if city not in city_to_coords:
                city_to_coords[city] = {'lat': row['lat'], 'lng': row['lng']}
        city_to_continent = geo_unique.set_index('city_ascii_upper')['continent'].to_dict()
        city_to_iso3 = geo_unique.set_index('city_ascii_upper')['iso3'].to_dict()
        
        # 2. CERDI data: create distance mapping
        distance_map = cerdi_df.set_index(['iso1', 'iso2'])['seadistance_km'].to_dict()
        
        # 3. EF data: create trade line mapping
        trade_map = ef_df.set_index(['Port of Origin', 'Port of Destination'])[
            ['Corresponding trade line group', 'Reefer', 'Dry']
        ].to_dict('index')
        
        # 4. Reefer data: create temperature control mapping
        reefer_df = reefer_df.copy()
        reefer_df['SpendCategory_Upper'] = reefer_df['SpendCategory'].str.strip().str.upper()
        temp_control_map = reefer_df.set_index('SpendCategory_Upper')['temp_control'].to_dict()
        
        # 5. Supplier data: create supplier mappings
        supplier_country_map = {}
        supplier_ef_map = {}
        for _, row in supplier_df.iterrows():
            if pd.notna(row['Supplier']):
                supplier_country_map[row['Supplier']] = row['country'] if pd.notna(row['country']) else None
                supplier_ef_map[row['Supplier']] = row['supplier_specific_ef_sc123'] if pd.notna(row['supplier_specific_ef_sc123']) else None
        
        return {
            'geo_df': geo_df,
            'cerdi_df': cerdi_df,
            'ef_df': ef_df,
            'reefer_df': reefer_df,
            'supplier_df': supplier_df,
            'precomputed_mappings': {
                'city_to_coords': city_to_coords,
                'city_to_continent': city_to_continent,
                'city_to_iso3': city_to_iso3,
                'distance_map': distance_map,
                'trade_map': trade_map,
                'temp_control_map': temp_control_map,
                'supplier_country_map': supplier_country_map,
                'supplier_ef_map': supplier_ef_map
            }
        }
        
    except Exception as e:
        st.error(f"Error processing geo_master.xlsx: {e}")
        import traceback
        st.error(f"Traceback: {traceback.format_exc()}")
        return None

def get_geo_master_data(geo_content, use_github=False):
    """Get processed geo master data with cached mappings, supporting both GitHub and local sources"""
    processed_data = load_and_process_geo_master(geo_content, use_github=use_github)
    if processed_data is None:
        return None, None, None, None, None, None
    
    return (
        processed_data['geo_df'],
        processed_data['cerdi_df'], 
        processed_data['ef_df'],
        processed_data['reefer_df'],
        processed_data['supplier_df'],
        processed_data['precomputed_mappings']
    )

def extract_spend_data(file_path):
    """Extracts spend data with expanded columns"""
    target_columns = [
        "YearReceipt", "PurchaseOrderNumber", "Supplier", 
        "ShippingMethod", "SpendCategory", "LineDescription",
        "FundType", "ShipToAddressCountry", "Freight per APO", 
        "Amount per PO Line"
    ]
    try:
        spend_df = pd.read_excel(
            file_path,
            usecols=target_columns
        )
        spend_df["SpendCategory"] = spend_df["SpendCategory"].str.strip().str.upper()
        spend_df["FundType"] = spend_df["FundType"].str.strip()
        return spend_df
    except Exception as e:
        st.error(f"Error processing spend data: {e}")
        return None

def enrich_bol_with_geo(bol_df, geo_df):
    """Enriches bol_df with geo information"""
    try:
        geo_unique = geo_df.drop_duplicates(subset=['city_ascii'], keep='first')
        geo_unique['city_ascii'] = geo_unique['city_ascii'].str.upper()
        city_to_iso3 = geo_unique.set_index('city_ascii')['iso3'].to_dict()
        city_to_continent = geo_unique.set_index('city_ascii')['continent'].to_dict()
        
        bol_df['Port of loading'] = bol_df['Port of loading'].str.upper()
        bol_df['Port of discharge'] = bol_df['Port of discharge'].str.upper()

        bol_df['Loading Port ISO3'] = bol_df['Port of loading'].map(city_to_iso3)
        bol_df['Loading Port Continent'] = bol_df['Port of loading'].map(city_to_continent)
        bol_df['Discharge Port ISO3'] = bol_df['Port of discharge'].map(city_to_iso3)
        bol_df['Discharge Port Continent'] = bol_df['Port of discharge'].map(city_to_continent)
        
        return bol_df
    except Exception as e:
        st.error(f"Error enriching bol data: {e}")
        return bol_df

def add_sea_distance(bol_df, cerdi_df):
    """Adds sea distance information"""
    try:
        distance_map = cerdi_df.set_index(['iso1', 'iso2'])['seadistance_km'].to_dict()
        bol_df['Sea Distance (km)'] = bol_df.apply(
            lambda row: distance_map.get(
                (row['Loading Port ISO3'], row['Discharge Port ISO3']), 
                None
            ),
            axis=1
        )
        return bol_df
    except Exception as e:
        st.error(f"Error adding sea distance: {e}")
        return bol_df

def add_trade_line_info(bol_df, ef_df):
    """Adds trade line information"""
    try:
        trade_map = ef_df.set_index(['Port of Origin', 'Port of Destination'])[
            ['Corresponding trade line group', 'Reefer', 'Dry']
        ].to_dict('index')
        
        bol_df[['Trade Line Group', 'Reefer', 'Dry']] = bol_df.apply(
            lambda row: pd.Series(
                trade_map.get(
                    (row['Loading Port Continent'], row['Discharge Port Continent']),
                    {'Corresponding trade line group': None, 'Reefer': None, 'Dry': None}
                )
            ),
            axis=1
        )
        return bol_df
    except Exception as e:
        st.error(f"Error adding trade line information: {e}")
        return bol_df

def add_spend_category(bol_df, spend_df):
    """Adds SpendCategory information"""
    try:
        spend_category_map = spend_df.drop_duplicates('PurchaseOrderNumber').set_index('PurchaseOrderNumber')['SpendCategory'].to_dict()
        bol_df['SpendCategory'] = bol_df['APO no.'].map(spend_category_map)
        return bol_df
    except Exception as e:
        st.error(f"Error adding spend category: {e}")
        return bol_df

def add_temp_control(bol_df, reefer_df):
    """Adds temperature control information with case-insensitive matching"""
    try:
        # Create case-insensitive mapping
        reefer_df = reefer_df.copy()
        reefer_df['SpendCategory_Upper'] = reefer_df['SpendCategory'].str.strip().str.upper()
        
        temp_control_map = reefer_df.drop_duplicates('SpendCategory_Upper').set_index('SpendCategory_Upper')['temp_control'].to_dict()
        
        # Convert bol_df categories to uppercase for matching
        if 'SpendCategory' in bol_df.columns:
            bol_df = bol_df.copy()
            bol_df['SpendCategory_Upper'] = bol_df['SpendCategory'].str.strip().str.upper()
            bol_df['Temperature Control'] = bol_df['SpendCategory_Upper'].map(temp_control_map)
            bol_df = bol_df.drop('SpendCategory_Upper', axis=1)
        
        return bol_df
    except Exception as e:
        st.error(f"Error adding temperature control: {e}")
        return bol_df

def add_teu_column(bol_df):
    """Adds TEU column"""
    try:
        bol_df = bol_df.copy()
        bol_df['teu'] = bol_df['Gross weight, kg'] / 10000
        return bol_df
    except Exception as e:
        st.error(f"Error calculating TEU: {e}")
        return bol_df

def add_ghg_emissions(bol_df):
    """Adds GHG emissions columns"""
    try:
        bol_df = bol_df.copy()
        bol_df['ghg_emissions_gCO2e'] = bol_df.apply(
            lambda row: (row['teu'] * row['Sea Distance (km)'] * 
                        (row['Reefer'] if row['Temperature Control'] == 'YES' else row['Dry'])),
            axis=1
        )
        bol_df['ghg_emissions_tCO2e'] = bol_df['ghg_emissions_gCO2e'] / 1000000
        return bol_df
    except Exception as e:
        st.error(f"Error calculating GHG emissions: {e}")
        return bol_df

def process_bol_file(bol_file, geo_content, spend_file=None, use_github=False):
    """Process BOL data with comprehensive error handling"""
    try:
        # Load geo master data with new function
        geo_data = get_geo_master_data(geo_content, use_github=use_github)
        if geo_data[0] is None:
            st.error("Failed to load geo master data")
            return None
            
        geo_df, cerdi_df, ef_df, reefer_df, supplier_df, precomputed_mappings = geo_data
        
        # Load BOL data
        bol_df = extract_bol_data(bol_file)
        if bol_df is None or bol_df.empty:
            st.error("No valid BOL data found or failed to load BOL file")
            return None
            
        required_columns = ['Gross weight, kg', 'Port of loading', 'Port of discharge', 'APO no.']
        missing_cols = [col for col in required_columns if col not in bol_df.columns]
        if missing_cols:
            st.error(f"BOL file is missing required columns: {', '.join(missing_cols)}")
            return None
            
        # Load spend data if available
        spend_df = None
        if spend_file is not None:
            spend_df = extract_spend_data(spend_file)
            if spend_df is None or spend_df.empty:
                st.warning("Unable to load spend data; BOL processing will continue without spend-dependent fields.")
                spend_df = None

        # Initialize processing
        final_bol_df = bol_df.copy()
        
        # Step 2: Data enrichment pipeline
        def safe_apply_step(df, step):
            try:
                missing = [col for col in step.get("required_cols", []) if col not in df.columns]
                if missing:
                    st.warning(f"Skipped {step['name']} - missing columns: {missing}")
                    return df
                result = step["func"](df.copy(), *step.get("deps", []))
                return result if result is not None else df
            except Exception as e:
                st.warning(f"Warning during {step['name']}: {str(e)}")
                return df

        enrichment_steps = [
            {"name": "Geo information", "func": enrich_bol_with_geo, "deps": [geo_df]},
            {"name": "Sea distance", "func": add_sea_distance, "deps": [cerdi_df], 
             "required_cols": ["Loading Port ISO3", "Discharge Port ISO3"]},
            {"name": "Trade line", "func": add_trade_line_info, "deps": [ef_df],
             "required_cols": ["Loading Port Continent", "Discharge Port Continent"]}
        ]
        if spend_df is not None:
            enrichment_steps += [
                {"name": "Spend category", "func": add_spend_category, "deps": [spend_df],
                 "required_cols": ["APO no."]},
                {"name": "Temperature", "func": add_temp_control, "deps": [reefer_df],
                 "required_cols": ["SpendCategory"]}
            ]

        for step in enrichment_steps:
            final_bol_df = safe_apply_step(final_bol_df, step)

        # Step 3: Calculations
        calculation_steps = [
            {
                "name": "TEU calculation",
                "func": lambda df: add_teu_column(df) if 'Gross weight, kg' in df.columns else df,
                "required_cols": ["Gross weight, kg"]
            },
            {
                "name": "GHG emissions",
                "func": lambda df: add_ghg_emissions(df) if all(col in df.columns for col in [
                    'teu', 'Sea Distance (km)', 'Temperature Control', 'Reefer', 'Dry'
                ]) else df,
                "required_cols": ["teu", "Sea Distance (km)", "Temperature Control", "Reefer", "Dry"]
            }
        ]

        for calc in calculation_steps:
            final_bol_df = safe_apply_step(final_bol_df, calc)

        # Final validation
        if final_bol_df.empty:
            st.error("Processing resulted in empty dataframe")
            return None
            
        if 'ghg_emissions_gCO2e' not in final_bol_df.columns:
            st.warning("GHG emissions calculation failed - missing required columns")
            
        return final_bol_df
        
    except Exception as e:
        st.error(f"Critical error processing BOL data: {str(e)}")
        return None

def calculate_scope31_emissions(spend_file, awb_data, bol_data, selected_years):
    """Calculate Scope 3.1 emissions with hardcoded EPA factors for all spend APOs."""
    try:
        # Hardcoded EPA data
        epa_data = [
            ["(PRO) Antiretroviral medicines", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antileishmaniasis", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antitubercular drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Disease vectors management and control", "Exterminating and Pest Control Services", "561700", 0.214],
            ["(PRO) Vaccines and antigens and toxoids", "Biological product (except diagnostics) manufacturing", "325414", 0.126],
            ["(PRO) Antibacterials", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Anti-malarial", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antineoplastic agents", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Cardiovascular Medicines", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Nutritional Supplements", "All Other Miscellaneous Food Manufacturing", "311990", 0.358],
            ["(PRO) Antiviral drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Clinical and diagnostic analyzers and accessories and supplies", "In-Vitro Diagnostic Substance Manufacturing", "325413", 0.161],
            ["(PRO) Anti-Chagas", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antifungal drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Syringes and accessories", "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Rapid test kits", "In-Vitro Diagnostic Substance Manufacturing", "325413", 0.161],
            ["(PRO) Antipsychotics", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Central nervous system drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Laboratory and scientific equipment", "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Needle or blade or sharps disposal container or cart", "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Veterinary vaccines and virology products", "Biological product (except diagnostics) manufacturing", "325414", 0.126],
            ["(PRO) Laboratory supplies and fixtures", "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Medical Equipment and Accessories and Supplies", "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(CC) (PAHO internal use) Equipment Services: installation, maintenance, leasing and training", "In-Vitro Diagnostic Substance Manufacturing", "325413", 0.161],
            ["(PRO) Hematolic drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Anaesthetic drugs and related adjuncts and analeptics",  "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Patient care and treatment products and supplies",  "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Antidotes and emetics", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Industrial freezers and refrigerators", "Air-Conditioning and Warm Air Heating Equipment and Commercial and Industrial Refrigeration Equipment Manufacturing", "333415", 0.156],
            ["(PRO) Antidiabetic agents and hyperglycemic agents", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Cold storage box", "Air-Conditioning and Warm Air Heating Equipment and Commercial and Industrial Refrigeration Equipment Manufacturing", "333415", 0.156],
            ["(PRO) Cold pack or ice brick",  "Air-Conditioning and Warm Air Heating Equipment and Commercial and Industrial Refrigeration Equipment Manufacturing", "333415", 0.156],
            ["(CC) (Technical Cooperation) Supplies and Materials",  "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Drugs affecting the respiratory tract",  "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Electrolytes",  "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO)Temperature and heat measuring instruments",  "Surgical and Medical Instrument Manufacturing", "339112", 0.119],
            ["(PRO) Muscle Relaxant Medicines",  "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(CC) (Technical Cooperation) IT Equipment and Accessories",  "Surgical and Medical Instrument Manufacturing", "339112", 0.119]
        ]
        # Additional entries added from updated PAHO EF_usaepa_exiobase.xlsx
        epa_data += [
            ["(PRO) Estrogens and progestins and internal contraceptives", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antihypertensive drugs", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Anticonvulsants", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antituberculosis medicines", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Hematologic medicines", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Corticosteroids", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Pharmaceutical refrigerator and freezer", "Air-Conditioning and Warm Air Heating Equipment and Commercial and Industrial Refrigeration Equipment Manufacturing", "333415", 0.156],
            ["(PRO) Indicators and Reagents", "In-Vitro Diagnostic Substance Manufacturing", "325413", 0.161],
            ["(PRO) Controlled substance analgesics", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antihistamines or H1 blockers", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Antiarrythmics and antianginals and cardioplegics and drugs for heart failure", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Anthelmintics", "Pharmaceutical Preparation Manufacturing", "325412", 0.045],
            ["(PRO) Chemical biological control equipment and accessories and supplies","Exterminating and Pest Control Services", "561700", 0.214],
            ["(PRO) Immunosuppressants", "Pharmaceutical Preparation Manufacturing", "325412", 0.045]
        ]
        
        epa_df = pd.DataFrame(epa_data, columns=['Spend Category', 'NAICS Title', 'NAICS Code', 'kg CO2e/USD'])
        
        # Load spend data
        spend_df = pd.read_excel(
            spend_file,
            usecols=[
                "YearReceipt", "PurchaseOrderNumber", "Supplier", 
                "ShippingMethod", "SpendCategory", "LineDescription",
                "FundType", "ShipToAddressCountry", "Freight per APO", 
                "Amount per PO Line"
            ]
        )
        
        # Clean and standardize spend categories
        spend_df['SpendCategory'] = spend_df['SpendCategory'].str.strip()
        
        # Keep all spend rows regardless of whether the APO appears in AWB/BOL
        filtered_spend = spend_df.copy()
        
        # Year filter
        if selected_years:
            filtered_spend = filtered_spend[
                filtered_spend['YearReceipt'].astype(str).isin(selected_years)
            ]
        
        # FundType filter
        filtered_spend = filtered_spend[
            filtered_spend['FundType'].isin(['Revolving Fund', 'Strategic Fund'])
        ]
        
        # Exclude donations
        filtered_spend = filtered_spend[
            ~filtered_spend['LineDescription'].str.contains('donations|discount', case=False, na=False)
        ]
        
        if filtered_spend.empty:
            st.warning("No spend data matches the selected filters.")
            return None
        
        # Merge using exact matching
        merged_df = pd.merge(
            filtered_spend,
            epa_df,
            left_on='SpendCategory',
            right_on='Spend Category',
            how='left'
        )
        
        # Debug: Show matching statistics
        total_rows = len(merged_df)
        matched_rows = len(merged_df[~merged_df['NAICS Code'].isna()])
        st.info(f"Matching rate: {matched_rows}/{total_rows} ({matched_rows/total_rows:.1%}) rows matched")
        
        if matched_rows == 0:
            st.error("No spend categories matched with EPA factors. Please check your data.")
            return None
        
        # Calculate emissions
        merged_df['kg CO2e'] = merged_df['Amount per PO Line'] * merged_df['kg CO2e/USD']
        merged_df['t CO2e'] = merged_df['kg CO2e'] / 1000
        
        # Final output columns
        result_df = merged_df[[
            'Supplier', 'SpendCategory', 'NAICS Title', 'NAICS Code',
            'Amount per PO Line', 'kg CO2e/USD', 'kg CO2e', 't CO2e',
            'YearReceipt', 'FundType', 'PurchaseOrderNumber'
        ]].rename(columns={
            'SpendCategory': 'Spend Category',
            'Amount per PO Line': 'Amount (USD)'
        })
        
        # Sort and add totals
        result_df = result_df.sort_values('Supplier', ascending=False)
        totals = pd.DataFrame([{
            'Supplier': 'TOTAL',
            'Amount (USD)': result_df['Amount (USD)'].sum(),
            'kg CO2e': result_df['kg CO2e'].sum(),
            't CO2e': result_df['t CO2e'].sum()
        }])
        
        return pd.concat([result_df, totals], ignore_index=True)
        
    except Exception as e:
        st.error(f"Error calculating Scope 3.1 emissions: {e}")
        return None

# ==================== EPA EF Functions ====================
@st.cache_data
def load_epa_ef():
    """Load EPA emission factors from file"""
    try:
        if USE_GITHUB_DATA:
            response = requests.get(GITHUB_EPA_EF_URL, timeout=10)
            response.raise_for_status()
            return response.content
        else:
            with open(LOCAL_EPA_EF_PATH, 'rb') as f:
                return f.read()
    except Exception as e:
        st.error(f"Failed to load EPA EF: {str(e)}")
        return None


# ==================== GEO MASTER LOADING ====================
@st.cache_data
def load_geo_master():
    """Load geo_master from GitHub or local with caching"""
    try:
        if USE_GITHUB_DATA:
            response = requests.get(GITHUB_GEO_MASTER_URL, timeout=30)
            response.raise_for_status()
            return response.content, True  # Return content + use_github flag
        else:
            with open(LOCAL_GEO_MASTER_PATH, 'rb') as f:
                return f.read(), False  # Return content + use_github flag
    except Exception as e:
        st.error(f"Failed to load geo_master: {str(e)}")
        return None, None

# ==================== EXIOBASE SCOPE 3.1 FUNCTION ====================
def calculate_scope31_emissions_exiobase(spend_file, geo_file, awb_data, bol_data, selected_years):
    """Calculate Scope 3.1 emissions using Exiobase emission factors for all spend APOs."""
    try:
        # Define the two sets of Exiobase emission factors
        # Group 1: Pharmaceutical/medical product categories (ID 24d)
        exiobase_group1 = {
            'AUSTRIA': 0.545131753,
            'BELGIUM': 0.479460088,
            'BRAZIL': 1.109644522,
            'BULGARIA': 1.339848257,
            'CANADA': 0.975906107,
            'CHINA': 2.331383703,
            'CYPRUS': 0.969797487,
            'DENMARK': 0.684560609,
            'FINLAND': 0.494782497,
            'FRANCE': 0.217503994,
            'GERMANY': 0.504585879,
            'GREECE': 1.063976939,
            'INDIA': 2.335286193,
            'INDONESIA': 1.249600634,
            'IRELAND': 0.479919211,
            'ITALY': 0.938398599,
            'NETHERLANDS': 0.976116293,
            'SOUTH KOREA': 1.455342401,
            'SPAIN': 0.753155613,
            'SWEDEN': 0.41669048,
            'SWITZERLAND': 0.641113126,
            'UNITED STATES': 0.515891942,
            'UNITED ARAB EMIRATES': 2.238255645,
            'ISRAEL': 2.238255645,
            'ARGENTINA': 2.605049361,
            'VIETNAM': 2.238255645,
            'PANAMA': 2.605049361,
            'JAPAN': 2.238255645,
            'COSTA RICA': 2.605049361,
            'COLOMBIA': 2.605049361,
            'URUGUAY': 2.605049361,
            'MALAYSIA': 2.238255645,
            'ANTIGUA AND BARBUDA': 2.605049361,
            'THAILAND': 2.238255645,
            'LUXEMBOURG': 0.52,
            'SOUTH AFRICA': 1.90,
            'UNITED KINGDOM': 0.29,
            'MEXICO': 1.09
        }

        # Group 2: Equipment/analyzers categories (ID 33)
        exiobase_group2 = {
            'AUSTRIA': 0.154586475,
            'BELGIUM': 0.440162225,
            'BRAZIL': 0.397086636,
            'BULGARIA': None,
            'CANADA': None,
            'CHINA': 0.980305411,
            'CYPRUS': None,
            'DENMARK': None,
            'FINLAND': None,
            'FRANCE': 0.266389395,
            'GERMANY': 0.180232507,
            'GREECE': None,
            'INDIA': 0.984197791,
            'INDONESIA': None,
            'IRELAND': 0.159060233,
            'ITALY': None,
            'NETHERLANDS': 0.292628813,
            'SOUTH KOREA': 0.474450783,
            'SPAIN': None,
            'SWEDEN': 0.25623846,
            'SWITZERLAND': 0.28,
            'UNITED STATES': 0.387415696,
            'UNITED ARAB EMIRATES': 1.236221092,
            'ISRAEL': 1.236221092,
            'ARGENTINA': 0.662505065,
            'VIETNAM': 1.236221092,
            'PANAMA': 0.662505065,
            'JAPAN': 1.236221092,
            'COSTA RICA': 0.662505065,
            'COLOMBIA': 0.662505065,
            'URUGUAY': 0.662505065,
            'MALAYSIA': 1.236221092,
            'ANTIGUA AND BARBUDA': 0.662505065,
            'THAILAND': 1.236221092,
            'SOUTH AFRICA': 1.02,
            'UNITED KINGDOM': 0.15,
            'MEXICO': 0.8
        }

        # Group 1 categories (ID 24d)
        group1_categories = [
            "(PRO) Antiretroviral medicines",
            "(PRO) Antileishmaniasis",
            "(PRO) Antitubercular drugs",
            "(PRO) Disease vectors management and control",
            "(PRO) Vaccines and antigens and toxoids",
            "(PRO) Antibacterials",
            "(PRO) Anti-malarial",
            "(PRO) Antineoplastic agents",
            "(PRO) Cardiovascular Medicines",
            "(PRO) Nutritional Supplements",
            "(PRO) Antiviral drugs",
            "(PRO) Anti-Chagas",
            "(PRO) Antifungal drugs",
            "(PRO) Syringes and accessories",
            "(PRO) Antipsychotics",
            "(PRO) Central nervous system drugs",
            "(PRO) Needle or blade or sharps disposal container or cart",
            "(PRO) Veterinary vaccines and virology products",
            "(PRO) Hematolic drugs",
            "(PRO) Anaesthetic drugs and related adjuncts and analeptics",
            "(PRO) Patient care and treatment products and supplies",
            "(PRO) Antidotes and emetics",
            "(PRO) Industrial freezers and refrigerators",
            "(PRO) Antidiabetic agents and hyperglycemic agents",
            "(PRO) Cold storage box",
            "(PRO) Cold pack or ice brick",
            "(CC) (Technical Cooperation) Supplies and Materials",
            "(PRO) Drugs affecting the respiratory tract",
            "(PRO) Electrolytes",
            "(PRO) Muscle Relaxant Medicines",
            "(CC) (Technical Cooperation) IT Equipment and Accessories"
        ]

        # Group 2 categories (ID 33)
        group2_categories = [
            "(PRO) Clinical and diagnostic analyzers and accessories and supplies",
            "(PRO) Rapid test kits",
            "(PRO) Laboratory and scientific equipment",
            "(PRO) Laboratory supplies and fixtures",
            "(PRO) Medical Equipment and Accessories and Supplies",
            "(CC) (PAHO internal use) Equipment Services: installation, maintenance, leasing and training",
            "(PRO)Temperature and heat measuring instruments"
        ]

        def standardize_country_name(country_input):
            """Convert any country format to standardized full name"""
            if pd.isna(country_input):
                return None
                
            country_str = str(country_input).upper().strip()
            
            # First check if it's already a standardized name we recognize
            if country_str in exiobase_group1 or country_str in exiobase_group2:
                return country_str
                
            # Try fuzzy matching for known countries
            known_countries = list(exiobase_group1.keys()) + list(exiobase_group2.keys())
            best_match, score = process.extractOne(country_str, known_countries)
            if score > 80:  # Good match threshold
                return best_match
                
            return None  # No match found

        # Load geo master data for supplier-country mapping
        if isinstance(geo_file, str):
            # It's a file path - read it
            with open(geo_file, 'rb') as f:
                geo_content = f.read()
            geo_data = get_geo_master_data(geo_content, use_github=False)
        else:
            # It's already content (BytesIO or bytes)
            geo_data = get_geo_master_data(geo_file, use_github=False)
        
        if geo_data[0] is None:
            st.error("Failed to load geo master data for supplier-country mapping")
            return None
            
        _, _, _, _, supplier_df, _ = geo_data
        
        # Create supplier to country mapping from the supplier worksheet (SAME APPROACH AS SUPPLIER-SPECIFIC METHOD)
        supplier_country_map = {}
        for _, row in supplier_df.iterrows():
            supplier = row['Supplier']
            country = row['country']
            if pd.notna(supplier) and pd.notna(country):
                standardized_country = standardize_country_name(country)
                if standardized_country:
                    supplier_country_map[supplier] = standardized_country
        
        # Load spend data
        spend_df = pd.read_excel(
            spend_file,
            usecols=[
                "YearReceipt", "PurchaseOrderNumber", "Supplier", 
                "ShippingMethod", "SpendCategory", "LineDescription",
                "FundType", "ShipToAddressCountry", "Freight per APO", 
                "Amount per PO Line"
            ]
        )
        
        # Clean and standardize spend categories
        spend_df['SpendCategory'] = spend_df['SpendCategory'].str.strip()
        
        # Keep all spend rows regardless of whether the APO appears in AWB/BOL
        filtered_spend = spend_df.copy()
        
        # Year filter
        if selected_years:
            filtered_spend = filtered_spend[
                filtered_spend['YearReceipt'].astype(str).isin(selected_years)
            ]
        
        # FundType filter
        filtered_spend = filtered_spend[
            filtered_spend['FundType'].isin(['Revolving Fund', 'Strategic Fund'])
        ]
        
        # Exclude donations
        filtered_spend = filtered_spend[
            ~filtered_spend['LineDescription'].str.contains('donations|discount', case=False, na=False)
        ]
        
        if filtered_spend.empty:
            st.warning("No spend data matches the selected filters.")
            return None

        # Add Exiobase calculations to the spend data
        result_df = filtered_spend.copy()
        
        # Determine which group each spend category belongs to and get appropriate EF
        result_df['Exiobase_Group'] = result_df['SpendCategory'].apply(
            lambda x: 'ID 24d' if x in group1_categories else ('ID 33' if x in group2_categories else None)
        )
        
        # Map supplier to standardized country name using the supplier worksheet mapping (SAME APPROACH)
        result_df['Supplier_Country'] = result_df['Supplier'].map(supplier_country_map)
        
        # Get Exiobase emission factor
        def get_exiobase_ef(group, country):
            if pd.isna(group) or pd.isna(country):
                return None
            if group == 'ID 24d':
                return exiobase_group1.get(country, None)
            elif group == 'ID 33':
                return exiobase_group2.get(country, None)
            return None
        
        result_df['Exiobase_kg_CO2e/USD'] = result_df.apply(
            lambda row: get_exiobase_ef(row['Exiobase_Group'], row['Supplier_Country']), axis=1
        )
        
        # Calculate emissions - CORRECTED: Exiobase factors are kg CO2e/USD, not g CO2e/USD
        result_df['Exiobase_kg_CO2e'] = result_df['Amount per PO Line'] * result_df['Exiobase_kg_CO2e/USD']
        result_df['Exiobase_t_CO2e'] = result_df['Exiobase_kg_CO2e'] / 1000  # Convert kg to tons
        
        # Final output columns
        final_columns = [
            'Supplier', 'SpendCategory', 'Supplier_Country', 'Exiobase_Group',
            'Amount per PO Line', 'Exiobase_kg_CO2e/USD', 'Exiobase_kg_CO2e', 'Exiobase_t_CO2e',
            'YearReceipt', 'FundType', 'PurchaseOrderNumber'
        ]
        
        result_df = result_df[final_columns].rename(columns={
            'SpendCategory': 'Spend Category',
            'Amount per PO Line': 'Amount (USD)',
            'Supplier_Country': 'Supplier Country'
        })
        
        # Sort and add totals
        result_df = result_df.sort_values('Supplier', ascending=False)
        totals = pd.DataFrame([{
            'Supplier': 'TOTAL',
            'Amount (USD)': result_df['Amount (USD)'].sum(),
            'Exiobase_kg_CO2e': result_df['Exiobase_kg_CO2e'].sum(),
            'Exiobase_t_CO2e': result_df['Exiobase_t_CO2e'].sum()
        }])
        
        return pd.concat([result_df, totals], ignore_index=True)
        
    except Exception as e:
        st.error(f"Error calculating Scope 3.1 emissions with Exiobase: {e}")
        return None
    
def calculate_scope31_emissions_supplier_specific(spend_file, geo_file, awb_data, bol_data, selected_years):
    """Calculate Scope 3.1 emissions using supplier-specific factors with Exiobase fallback for all spend APOs."""
    try:
        # Load geo master data for supplier-specific factors and country mapping
        if isinstance(geo_file, str):
            # It's a file path - read it
            with open(geo_file, 'rb') as f:
                geo_content = f.read()
            geo_data = get_geo_master_data(geo_content, use_github=False)
        else:
            # It's already content (BytesIO or bytes)
            geo_data = get_geo_master_data(geo_file, use_github=False)        
        
        if geo_data[0] is None:
            st.error("Failed to load geo master data for supplier-specific calculation")
            return None
        geo_df, _, _, _, supplier_df, precomputed_mappings = geo_data
        
        # Create supplier mapping with both country and supplier-specific EF
        supplier_info_map = {}
        for _, row in supplier_df.iterrows():
            supplier = row['Supplier']
            if pd.notna(supplier):
                supplier_info_map[supplier] = {
                    'country': row['country'] if pd.notna(row['country']) else None,
                    'supplier_specific_ef': row['supplier_specific_ef_sc123'] if pd.notna(row['supplier_specific_ef_sc123']) else None
                }
        
        # Load spend data
        spend_df = pd.read_excel(
            spend_file,
            usecols=[
                "YearReceipt", "PurchaseOrderNumber", "Supplier", 
                "ShippingMethod", "SpendCategory", "LineDescription",
                "FundType", "ShipToAddressCountry", "Freight per APO", 
                "Amount per PO Line"
            ]
        )
        
        # Clean and standardize spend categories
        spend_df['SpendCategory'] = spend_df['SpendCategory'].str.strip()
        
        # Keep all spend rows regardless of whether the APO appears in AWB/BOL
        filtered_spend = spend_df.copy()
        
        # Year filter
        if selected_years:
            filtered_spend = filtered_spend[
                filtered_spend['YearReceipt'].astype(str).isin(selected_years)
            ]
        
        # FundType filter
        filtered_spend = filtered_spend[
            filtered_spend['FundType'].isin(['Revolving Fund', 'Strategic Fund'])
        ]
        
        # Exclude donations
        filtered_spend = filtered_spend[
            ~filtered_spend['LineDescription'].str.contains('donations|discount', case=False, na=False)
        ]
        
        if filtered_spend.empty:
            st.warning("No spend data matches the selected filters.")
            return None
        
        # Add supplier information to spend data
        result_df = filtered_spend.copy()
        
        # Extract country and supplier-specific EF directly - FIXED: Handle missing suppliers
        result_df['Supplier_Country'] = result_df['Supplier'].apply(
            lambda x: supplier_info_map.get(x, {}).get('country') if pd.notna(x) else None
        )
        result_df['Supplier_Specific_EF'] = result_df['Supplier'].apply(
            lambda x: supplier_info_map.get(x, {}).get('supplier_specific_ef') if pd.notna(x) else None
        )
        
        # Determine calculation method for each row
        result_df['Calculation_Method'] = result_df['Supplier_Specific_EF'].apply(
            lambda x: 'Supplier-Specific' if pd.notna(x) else 'Exiobase'
        )
        
        # Show method distribution
        method_counts = result_df['Calculation_Method'].value_counts()
        st.info(f"Calculation method distribution:\n{method_counts}")
        
        # Get Exiobase factors for fallback (reuse existing function)
        exiobase_result = calculate_scope31_emissions_exiobase(
            spend_file, geo_file, awb_data, bol_data, selected_years
        )
        
        if exiobase_result is None:
            st.error("Failed to get Exiobase factors for fallback calculation")
            return None
        
        # Create a dictionary for Exiobase factors instead of merging
        # This prevents the cartesian product issue with duplicate suppliers
        exiobase_dict = {}
        for _, row in exiobase_result.iterrows():
            if row['Supplier'] != 'TOTAL':  # Skip the total row
                key = (row['Supplier'], row['Spend Category'])
                exiobase_dict[key] = row['Exiobase_kg_CO2e/USD']
        
        # Map Exiobase factors using the dictionary - much more efficient
        result_df['Exiobase_kg_CO2e/USD'] = result_df.apply(
            lambda row: exiobase_dict.get((row['Supplier'], row['SpendCategory']), None),
            axis=1
        )
        
        # Calculate emissions based on method - SIMPLIFIED and SAFE
        def calculate_emissions(row):
            try:
                amount = row['Amount per PO Line']
                if pd.isna(amount) or amount == 0:
                    return 0
                
                if (row['Calculation_Method'] == 'Supplier-Specific' and 
                    pd.notna(row['Supplier_Specific_EF'])):
                    # Use supplier-specific factor
                    ef = row['Supplier_Specific_EF']
                    if ef < 0 or ef > 1000:  # Sanity check: EF should be reasonable
                        st.warning(f"Unusual supplier-specific EF for {row['Supplier']}: {ef}")
                        return 0
                    return ef * amount
                
                elif pd.notna(row['Exiobase_kg_CO2e/USD']):
                    # Use Exiobase factor
                    ef = row['Exiobase_kg_CO2e/USD']
                    if ef < 0 or ef > 1000:  # Sanity check
                        st.warning(f"Unusual Exiobase EF for {row['Supplier']}: {ef}")
                        return 0
                    return ef * amount
                
                else:
                    return 0  # No valid factors
                    
            except Exception as e:
                # Silent error handling to avoid breaking the entire calculation
                return 0
        
        result_df['kg_CO2e'] = result_df.apply(calculate_emissions, axis=1)
        result_df['t_CO2e'] = result_df['kg_CO2e'] / 1000
        
        # Check for astronomical values
        total_emissions = result_df['kg_CO2e'].sum()
        max_emission = result_df['kg_CO2e'].max()
        
        st.info(f"Total emissions: {total_emissions:,.2f} kg CO2e")
        st.info(f"Maximum row emission: {max_emission:,.2f} kg CO2e")
        
        if total_emissions > 1e9:  # More than 1 billion kg CO2e
            st.warning("⚠️ Very large total emissions detected. Checking data...")
            
            # Check supplier-specific factors
            supplier_specific_ef_stats = result_df[result_df['Calculation_Method'] == 'Supplier-Specific']['Supplier_Specific_EF'].describe()
            st.warning(f"Supplier-specific EF stats:\n{supplier_specific_ef_stats}")
            
            # Check spend amounts
            amount_stats = result_df['Amount per PO Line'].describe()
            st.warning(f"Spend amount stats:\n{amount_stats}")
            
            # Show top 10 emitting rows
            top_emitters = result_df.nlargest(10, 'kg_CO2e')[['Supplier', 'Amount per PO Line', 'Supplier_Specific_EF', 'Exiobase_kg_CO2e/USD', 'kg_CO2e']]
            st.warning(f"Top 10 emitting rows:\n{top_emitters}")
        
        # Final output columns
        final_columns = [
            'Supplier', 'SpendCategory', 'Supplier_Country', 'Calculation_Method',
            'Supplier_Specific_EF', 'Exiobase_kg_CO2e/USD', 'Amount per PO Line', 
            'kg_CO2e', 't_CO2e', 'YearReceipt', 'FundType', 'PurchaseOrderNumber'
        ]
        
        result_df = result_df[final_columns].rename(columns={
            'SpendCategory': 'Spend Category',
            'Amount per PO Line': 'Amount (USD)',
            'Supplier_Country': 'Supplier Country',
            'Supplier_Specific_EF': 'Supplier Specific EF (kg CO2e/USD)',
            'Exiobase_kg_CO2e/USD': 'Exiobase EF (kg CO2e/USD)'
        })
        
        # Sort and add totals
        result_df = result_df.sort_values('Supplier', ascending=False)
        
        # Create totals row without including it in the main dataframe
        totals = pd.DataFrame([{
            'Supplier': 'TOTAL',
            'Spend Category': '',
            'Supplier Country': '',
            'Calculation_Method': '',
            'Supplier Specific EF (kg CO2e/USD)': '',
            'Exiobase EF (kg CO2e/USD)': '',
            'Amount (USD)': result_df['Amount (USD)'].sum(),
            'kg_CO2e': result_df['kg_CO2e'].sum(),
            't_CO2e': result_df['t_CO2e'].sum(),
            'YearReceipt': '',
            'FundType': '',
            'PurchaseOrderNumber': ''
        }])
        
        return pd.concat([result_df, totals], ignore_index=True)
        
    except Exception as e:
        st.error(f"Error calculating Scope 3.1 emissions with Supplier-Specific + Exiobase: {e}")
        import traceback
        st.error(f"Traceback: {traceback.format_exc()}")
        return None

def get_top_flight_routes(awb_data, n=5):
    """Get top flight routes by count (number of flights)"""
    if awb_data is None or awb_data.empty:
        return pd.DataFrame()
    
    # Filter out TOTAL rows and get only flight legs
    flight_data = awb_data[awb_data['leg'] != 'TOTAL'].copy()
    
    if flight_data.empty:
        return pd.DataFrame()
    
    # Group by origin-destination pairs and count flights
    route_counts = flight_data.groupby(['origin', 'destination']).agg({
        'leg': 'count',
        'ghg_emissions_tCO2e': 'sum',
        'distance_km': 'mean'
    }).rename(columns={
        'leg': 'flight_count',
        'distance_km': 'avg_distance_km'
    }).reset_index()
    
    # Create route label
    route_counts['route'] = route_counts['origin'] + ' → ' + route_counts['destination']
    
    # Sort by flight count and get top N
    top_routes = route_counts.nlargest(n, 'flight_count')
    
    return top_routes[['route', 'flight_count', 'ghg_emissions_tCO2e', 'avg_distance_km']]

def get_top_source_locations(awb_data, n=5):
    """Get top true source locations by count (number of shipments) - only first origins from AWB data only"""
    source_counts = []
    
    # Process AWB data only - only get the FIRST origin for each AWB
    if awb_data is not None and not awb_data.empty:
        # Get the first leg (lowest leg number) for each AWB to find the true origin
        first_origins = awb_data[awb_data['leg'] != 'TOTAL'].copy()
        first_origins = first_origins.sort_values(['awb_row', 'leg']).groupby('awb_row').first().reset_index()
        
        awb_sources = first_origins.groupby('origin').agg({
            'awb_row': 'count',
            'ghg_emissions_tCO2e': 'sum'
        }).rename(columns={
            'awb_row': 'shipment_count'
        }).reset_index()
        awb_sources['transport_type'] = 'Air'
        source_counts.append(awb_sources.rename(columns={'origin': 'location'}))
    
    if not source_counts:
        return pd.DataFrame()
    
    # Combine and get top locations by shipment count
    all_sources = pd.concat(source_counts, ignore_index=True)
    top_locations = all_sources.groupby('location').agg({
        'shipment_count': 'sum',
        'ghg_emissions_tCO2e': 'sum'
    }).reset_index()
    
    # Get transport type composition for each location
    transport_composition = all_sources.groupby(['location', 'transport_type'])['shipment_count'].sum().unstack(fill_value=0)
    transport_composition = transport_composition.reset_index()
    
    # Merge composition data
    top_locations = top_locations.merge(transport_composition, on='location', how='left')
    
    # Sort by shipment count and get top N
    top_locations = top_locations.nlargest(n, 'shipment_count')
    
    return top_locations

# ==================== TCE 1 CALCULATION FUNCTIONS ====================
def get_continent_from_coords(lat, lon):
    """Get continent from latitude and longitude coordinates"""
    if pd.isna(lat) or pd.isna(lon):
        return None
    
    # Simple continent mapping based on coordinates
    # This is a simplified approach - you might want to use a more robust geocoding service
    if -170 <= lon <= -30:  # Americas
        if 10 <= lat <= 80:  # North America
            return "North America"
        elif -60 <= lat <= 15:  # South America
            return "South America"
    elif -20 <= lon <= 60:  # Africa/Europe
        if -40 <= lat <= 40:  # Africa
            return "Africa"
        elif 35 <= lat <= 75:  # Europe
            return "Europe"
    elif 60 <= lon <= 180:  # Asia/Australia
        if -15 <= lat <= 60:  # Asia
            return "Asia"
    
    return None

def get_distance_by_continent(continent):
    """Get distance based on continent using the provided values"""
    distance_map = {
        "Asia": 662.00,
        "Europe": 100.00,
        "North America": 330.00,
        "South America": 600.00,
        "Africa": 618.31
    }
    return distance_map.get(continent, 0.0)

def calculate_tce1_emissions(weight_kg, distance_km):
    """Calculate TCE 1 emissions using the provided emission factor"""
    if pd.isna(weight_kg) or weight_kg <= 0 or pd.isna(distance_km) or distance_km <= 0:
        return 0.0
    
    # Convert kg to tons
    weight_ton = weight_kg / 1000
    
    # Emission factor: 1102.28 g CO2e/ton-km
    ef = 1102.28  # Updated to 1102.28 g CO2e/ton-km based on the latest information provided
    
    # Calculate emissions in g CO2e
    emissions_g = weight_ton * distance_km * ef
    
    # Convert to t CO2e
    emissions_t = emissions_g / 1000000
    
    return emissions_t

def calculate_tce1_breakdown(awb_data, bol_data, precomputed_mappings):
    """Calculate TCE 1 emissions for both air and sea shipments"""
    tce1_results = {'Air': 0.0, 'Sea': 0.0, 'Total': 0.0}
    
    # Process AWB data (Air shipments)
    if awb_data is not None and not awb_data.empty:
        # Get only the first leg of each AWB to get the true origin
        first_legs = awb_data[awb_data['leg'] != 'TOTAL'].copy()
        first_legs = first_legs.sort_values(['awb_row', 'leg']).groupby('awb_row').first().reset_index()
        
        for _, row in first_legs.iterrows():
            if pd.notna(row['gross_weight']) and row['gross_weight'] > 0:
                # Get continent from coordinates
                continent = None
                if pd.notna(row['origin_lat']) and pd.notna(row['origin_lon']):
                    continent = get_continent_from_coords(row['origin_lat'], row['origin_lon'])
                
                # If continent not found from coordinates, try to get from airport code
                if continent is None and pd.notna(row['origin']):
                    try:
                        origin_coords = get_airport_coords(row['origin'])
                        if None not in origin_coords:
                            continent = get_continent_from_coords(origin_coords[0], origin_coords[1])
                    except:
                        pass
                
                # Get distance based on continent
                distance_km = get_distance_by_continent(continent) if continent else 0.0
                
                # Calculate TCE 1 emissions
                tce1_emissions = calculate_tce1_emissions(row['gross_weight'], distance_km)
                
                tce1_results['Air'] += tce1_emissions
                tce1_results['Total'] += tce1_emissions
    
    # Process BOL data (Sea shipments)
    if bol_data is not None and not bol_data.empty:
        # Use precomputed city_to_continent mapping
        city_to_continent = precomputed_mappings['city_to_continent']
        
        for _, row in bol_data.iterrows():
            if pd.notna(row['Gross weight, kg']) and row['Gross weight, kg'] > 0:
                # Get continent from BOL data using precomputed mapping
                continent = None
                if 'Loading Port Continent' in row and pd.notna(row['Loading Port Continent']):
                    continent = row['Loading Port Continent']
                elif pd.notna(row['Port of loading']):
                    port_upper = row['Port of loading'].upper()
                    continent = city_to_continent.get(port_upper)
                
                # Get distance based on continent
                distance_km = get_distance_by_continent(continent) if continent else 0.0
                
                # Calculate TCE 1 emissions
                tce1_emissions = calculate_tce1_emissions(row['Gross weight, kg'], distance_km)
                
                tce1_results['Sea'] += tce1_emissions
                tce1_results['Total'] += tce1_emissions
    
    return tce1_results

# ==================== TCE 2 & 4 CALCULATION FUNCTIONS ====================
def calculate_tce_emissions(weight_kg, transport_mode, cargo_type):
    """
    Calculate TCE 2 or TCE 4 emissions based on transport mode and cargo type
    
    Parameters:
    - weight_kg: Weight in kilograms
    - transport_mode: 'Air' or 'Sea'
    - cargo_type: 'Reefer' or 'Dry'
    """
    if pd.isna(weight_kg) or weight_kg <= 0:
        return 0.0
    
    # Convert kg to tons
    weight_ton = weight_kg / 1000
    
    # Define emission factors (g CO2e/ton)
    if transport_mode == 'Air':
        if cargo_type == 'Reefer':
            ef = 2600  # g CO2e/ton for air reefer
        else:
            ef = 1200  # g CO2e/ton for air dry
    else:  # Sea
        if cargo_type == 'Reefer':
            ef = 1340  # g CO2e/ton for sea reefer
        else:
            ef = 1140  # g CO2e/ton for sea dry
    
    # Calculate emissions in g CO2e
    emissions_g = weight_ton * ef
    
    # Convert to t CO2e
    emissions_t = emissions_g / 1000000
    
    return emissions_t

def calculate_tce2_breakdown(awb_data, bol_data):
    """Calculate TCE 2 emissions (export hub operations) for both air and sea shipments"""
    tce2_results = {'Air': 0.0, 'Sea': 0.0, 'Total': 0.0}
    
    # Process AWB data (Air shipments)
    if awb_data is not None and not awb_data.empty:
        # Get only the first leg of each AWB to get the true origin data
        first_legs = awb_data[awb_data['leg'] != 'TOTAL'].copy()
        first_legs = first_legs.sort_values(['awb_row', 'leg']).groupby('awb_row').first().reset_index()
        
        for _, row in first_legs.iterrows():
            if pd.notna(row['gross_weight']) and row['gross_weight'] > 0:
                # Determine cargo type based on temperature control
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                
                # Calculate TCE 2 emissions
                tce2_emissions = calculate_tce_emissions(row['gross_weight'], 'Air', cargo_type)
                
                tce2_results['Air'] += tce2_emissions
                tce2_results['Total'] += tce2_emissions
    
    # Process BOL data (Sea shipments)
    if bol_data is not None and not bol_data.empty:
        for _, row in bol_data.iterrows():
            if pd.notna(row['Gross weight, kg']) and row['Gross weight, kg'] > 0:
                # Determine cargo type based on temperature control
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                
                # Calculate TCE 2 emissions
                tce2_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                
                tce2_results['Sea'] += tce2_emissions
                tce2_results['Total'] += tce2_emissions
    
    return tce2_results

def calculate_tce4_breakdown(awb_data, bol_data):
    """Calculate TCE 4 emissions (import hub operations) for both air and sea shipments"""
    tce4_results = {'Air': 0.0, 'Sea': 0.0, 'Total': 0.0}
    
    # Process AWB data (Air shipments)
    if awb_data is not None and not awb_data.empty:
        # Get only the first leg of each AWB to get the true origin data
        first_legs = awb_data[awb_data['leg'] != 'TOTAL'].copy()
        first_legs = first_legs.sort_values(['awb_row', 'leg']).groupby('awb_row').first().reset_index()
        
        for _, row in first_legs.iterrows():
            if pd.notna(row['gross_weight']) and row['gross_weight'] > 0:
                # Determine cargo type based on temperature control
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                
                # Calculate TCE 4 emissions
                tce4_emissions = calculate_tce_emissions(row['gross_weight'], 'Air', cargo_type)
                
                tce4_results['Air'] += tce4_emissions
                tce4_results['Total'] += tce4_emissions
    
    # Process BOL data (Sea shipments)
    if bol_data is not None and not bol_data.empty:
        for _, row in bol_data.iterrows():
            if pd.notna(row['Gross weight, kg']) and row['Gross weight, kg'] > 0:
                # Determine cargo type based on temperature control
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                
                # Calculate TCE 4 emissions
                tce4_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                
                tce4_results['Sea'] += tce4_emissions
                tce4_results['Total'] += tce4_emissions
    
    return tce4_results

# ==================== INTENSITY CALCULATION FUNCTIONS ====================
def calculate_scope31_intensity(scope31_results, spend_file, selected_years, method_type="epa"):
    """Calculate Scope 3.1 intensity (g CO2e/$) using FILTERED spend data."""
    if scope31_results is None or scope31_results.empty:
        return None

    # Get total emissions in grams based on method type
    total_emissions_g = 0
    if method_type == "epa":
        if 'kg CO2e' in scope31_results.columns:
            total_row = scope31_results[scope31_results['Supplier'] == 'TOTAL']
            if not total_row.empty:
                total_kg = total_row['kg CO2e'].values[0]
                total_emissions_g = total_kg * 1000  # Convert kg to grams
    elif method_type == "exiobase":
        if 'Exiobase_kg_CO2e' in scope31_results.columns:
            total_row = scope31_results[scope31_results['Supplier'] == 'TOTAL']
            if not total_row.empty:
                total_kg = total_row['Exiobase_kg_CO2e'].values[0]
                total_emissions_g = total_kg * 1000
    elif method_type == "supplier":
        if 'kg_CO2e' in scope31_results.columns:
            total_row = scope31_results[scope31_results['Supplier'] == 'TOTAL']
            if not total_row.empty:
                total_kg = total_row['kg_CO2e'].values[0]
                total_emissions_g = total_kg * 1000

    # Load spend data with FILTERS applied
    try:
        spend_df = pd.read_excel(
            spend_file,
            usecols=["YearReceipt", "FundType", "LineDescription", "Amount per PO Line"]
        )
        
        # Apply filters (same logic as Scope 3.1 calculation)
        filtered_spend = spend_df.copy()
        
        # Year filter
        if selected_years:
            filtered_spend = filtered_spend[
                filtered_spend['YearReceipt'].astype(str).isin(selected_years)
            ]
        
        # FundType filter (Strategic/Revolving only)
        filtered_spend = filtered_spend[
            filtered_spend['FundType'].isin(['Revolving Fund', 'Strategic Fund'])
        ]
        
        # Exclude donations/discounts (case-insensitive)
        filtered_spend = filtered_spend[
            ~filtered_spend['LineDescription'].str.contains('donations|discount', case=False, na=False)
        ]
        
        total_spend = filtered_spend['Amount per PO Line'].sum()
        
        if total_spend > 0 and total_emissions_g > 0:
            return total_emissions_g / total_spend
        else:
            return None
    except Exception as e:
        print(f"Error calculating filtered Scope 3.1 intensity: {e}")
        return None

def calculate_scope34_intensity(awb_data, bol_data):
    """Calculate Scope 3.4 intensity (g CO2e/ton-km)"""
    total_emissions_g = 0
    total_ton_km = 0
    
    # Process AWB data
    if awb_data is not None and not awb_data.empty:
        # Filter out TOTAL rows
        awb_legs = awb_data[awb_data['leg'] != 'TOTAL']
        
        # Sum emissions (convert from tCO2e to gCO2e) - coerce to numeric for safety
        if 'ghg_emissions_tCO2e' in awb_legs.columns:
            awb_emissions_t = pd.to_numeric(awb_legs['ghg_emissions_tCO2e'], errors='coerce').sum()
            if pd.isna(awb_emissions_t):
                awb_emissions_t = 0.0
            total_emissions_g += awb_emissions_t * 1_000_000
        
        # Calculate ton-km for air shipments
        for _, row in awb_legs.iterrows():
            if (pd.notna(row['gross_weight_ton']) and row['gross_weight_ton'] > 0 and 
                pd.notna(row['distance_km']) and row['distance_km'] > 0):
                total_ton_km += row['gross_weight_ton'] * row['distance_km']
    
    # Process BOL data
    if bol_data is not None and not bol_data.empty:
        # Sum emissions (convert from tCO2e to gCO2e) - coerce to numeric for safety
        if 'ghg_emissions_tCO2e' in bol_data.columns:
            bol_emissions_t = pd.to_numeric(bol_data['ghg_emissions_tCO2e'], errors='coerce').sum()
            if pd.isna(bol_emissions_t):
                bol_emissions_t = 0.0
            total_emissions_g += bol_emissions_t * 1_000_000

        # Calculate ton-km for ocean shipments
        for _, row in bol_data.iterrows():
            gw = pd.to_numeric(row.get('Gross weight, kg', None), errors='coerce')
            sd = pd.to_numeric(row.get('Sea Distance (km)', None), errors='coerce')
            if pd.notna(gw) and gw > 0 and pd.notna(sd) and sd > 0:
                weight_ton = gw / 1000.0  # Convert kg to tons
                total_ton_km += float(weight_ton) * float(sd)
    
    # Calculate intensity
    if total_ton_km > 0:
        intensity = total_emissions_g / total_ton_km
        return intensity
    else:
        return None


# ==================== UPDATED MAIN FUNCTION ====================
def main():
    st.title("✈️ GHG Emissions Calculator")
    st.markdown("Calculate emissions from Air Waybills (AWB) and Bills of Lading (BOL)")

    # ==================== SIDEBAR FILTERS ====================
    st.sidebar.header("Filters")
    
    # Analysis scope selection - FIXED KEY
    analysis_scope = st.sidebar.radio(
        "Select Scope to Analyze:",
        options=["Scope 3.1 (Purchased Goods)", "Scope 3.4 (Transportation)", "Both"],
        index=2,
        key="main_scope_selector"
    )
    
    # Single year selection - FIXED KEY
    # Build year options from 2022 up to the current year (inclusive)
    current_year = datetime.datetime.now().year
    # Present most-recent-first for convenience
    years = [str(y) for y in range(current_year, 2021, -1)]
    # Default to 2025 when available, otherwise fall back to the most recent year
    default_year = "2025"
    try:
        default_index = years.index(default_year)
    except ValueError:
        default_index = 0

    selected_year = st.sidebar.selectbox(
        "Select Year",
        options=years,
        index=default_index,
        key="year_filter_selectbox"
    )

    # Initialize empty lists for supplier and APO filters
    selected_suppliers = []
    selected_apos = []

    # ==================== FILE UPLOADS ====================
    with st.expander("📁 Upload Required Files", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            main_file = st.file_uploader(
                "1. Main Data (with AWB/BoL worksheets)",
                type=['xlsx'],
                key="main_data_uploader"
            )
            
        with col2:
            spend_file = st.file_uploader(
                "2. Spend Data (PAHO Spend DATA 22-24.xlsx)",
                type=['xlsx'],
                key="spend_data_uploader"
            )

    if not main_file:
        st.info("Please upload the main data file to begin")
        st.stop()

    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"] and not spend_file:
        st.error("Please upload the spend data file for Scope 3.1 analysis")
        st.stop()

    # ==================== GEO MASTER LOADING ====================
    geo_content, use_github = load_geo_master()
    if geo_content is None:
        st.stop()

    # Load geo master with precomputed mappings once
    geo_data = get_geo_master_data(geo_content, use_github=use_github)
    if geo_data[0] is None:
        st.error("Failed to load geo master data")
        st.stop()

    geo_df, cerdi_df, ef_df, reefer_df, supplier_df, precomputed_mappings = geo_data

    # ==================== EPA EF LOADING ====================
    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"]:
        epa_content = load_epa_ef()
        if epa_content is None:
            st.error("EPA emission factors are required for Scope 3.1 calculation")
            st.stop()
    else:
        epa_content = None

    # ==================== FILE PROCESSING ====================
    try:
        xls = pd.ExcelFile(main_file)
        sheet_names = xls.sheet_names
        has_awb = 'AWB' in sheet_names
        has_bol = 'BoL' in sheet_names
        
        if not has_awb and not has_bol:
            st.error("File must contain either 'AWB' or 'BoL' worksheet")
            st.stop()
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        st.stop()

    # Initialize variables to avoid UnboundLocalError
    bol_with_tce = None
    awb_with_tce = None

    # Process Scope 3.4 data first
    awb_results = None
    bol_results = None

    if has_awb:
        with st.spinner("Processing AWB data..."):
            awb_results = process_awb_file(main_file, sheet_name='AWB', 
                                        spend_file=spend_file, geo_content=geo_content, use_github=use_github)
            # Initialize awb_with_tce if we have AWB data
            if awb_results is not None:
                awb_with_tce = awb_results.copy()

    if has_bol:
        with st.spinner("Processing BOL data..."):
            bol_results = process_bol_file(main_file, geo_content, spend_file if spend_file is not None else None, use_github=use_github)
            if bol_results is not None:
                bol_with_tce = bol_results.copy()
                
                # ===== ADD TCE COLUMNS AND CALCULATIONS HERE =====
                # Initialize TCE columns
                bol_with_tce['TCE1_Emissions_tCO2e'] = 0.0
                bol_with_tce['TCE2_Emissions_tCO2e'] = 0.0
                bol_with_tce['TCE4_Emissions_tCO2e'] = 0.0
                
                # Use precomputed city_to_continent mapping instead of loading geo data again
                city_to_continent = precomputed_mappings['city_to_continent']
                
                # Calculate TCE emissions for each BOL row
                for idx, row in bol_with_tce.iterrows():
                    if pd.notna(row['Gross weight, kg']) and row['Gross weight, kg'] > 0:
                        # TCE 1: Inland transport to export hub
                        continent = None
                        if 'Loading Port Continent' in row and pd.notna(row['Loading Port Continent']):
                            continent = row['Loading Port Continent']
                        elif pd.notna(row['Port of loading']):
                            port_upper = row['Port of loading'].upper()
                            continent = city_to_continent.get(port_upper)
                        
                        distance_km = get_distance_by_continent(continent) if continent else 0.0
                        tce1_emissions = calculate_tce1_emissions(row['Gross weight, kg'], distance_km)
                        bol_with_tce.at[idx, 'TCE1_Emissions_tCO2e'] = tce1_emissions
                        
                        # TCE 2: Export hub operations
                        cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                        tce2_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                        bol_with_tce.at[idx, 'TCE2_Emissions_tCO2e'] = tce2_emissions
                        
                        # TCE 4: Import hub operations (same calculation as TCE 2 for simplicity)
                        tce4_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                        bol_with_tce.at[idx, 'TCE4_Emissions_tCO2e'] = tce4_emissions

    # ==================== UPDATE FILTERS ====================
    if awb_results is not None or bol_results is not None:
        # Get unique suppliers and APOs from both datasets
        all_suppliers = set()
        all_apos = set()
        
        if awb_results is not None and not awb_results.empty:
            all_suppliers.update(awb_results['Supplier'].dropna().unique())
            all_apos.update(awb_results['APO'].dropna().unique())
            
        if bol_results is not None and not bol_results.empty:
            all_suppliers.update(bol_results['Shipper name'].dropna().unique())
            all_apos.update(bol_results['APO no.'].dropna().unique())
        
        # Convert to sorted lists for the select boxes
        supplier_list = sorted([s for s in all_suppliers if pd.notna(s)])
        apo_list = sorted([str(a) for a in all_apos if pd.notna(a)])
        
        # Create sidebar filters
        selected_suppliers = st.sidebar.multiselect(
            "Filter by Supplier",
            options=supplier_list,
            default=[],
            key="supplier_filter"
        )
        
        selected_apos = st.sidebar.multiselect(
            "Filter by APO no.",
            options=apo_list,
            default=[],
            key="apo_filter"
        )

    # ==================== FILTER DATA ====================
    def filter_data(df, is_awb=True):
        if df is None or df.empty:
            return df
            
        filtered = df.copy()
        
        # Apply supplier filter
        if selected_suppliers:
            if is_awb:
                filtered = filtered[filtered['Supplier'].isin(selected_suppliers)]
            else:
                filtered = filtered[filtered['Shipper name'].isin(selected_suppliers)]
        
        # Apply APO filter
        if selected_apos:
            if is_awb:
                filtered = filtered[filtered['APO'].astype(str).isin(selected_apos)]
            else:
                filtered = filtered[filtered['APO no.'].astype(str).isin(selected_apos)]
        
        return filtered

    filtered_awb = filter_data(awb_results, is_awb=True)
    filtered_bol = filter_data(bol_results, is_awb=False)

    # ==================== SCOPE 3.1 PROCESSING ====================
    # NOW process Scope 3.1 data AFTER filtering
    scope31_results = None
    scope31_exiobase_results = None
    scope31_supplier_specific_results = None

    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"] and spend_file and epa_content:
        with st.spinner("Processing Scope 3.1 data..."):
            scope31_results = calculate_scope31_emissions(spend_file, filtered_awb, filtered_bol, [selected_year])
            scope31_exiobase_results = calculate_scope31_emissions_exiobase(
                spend_file, geo_content, filtered_awb, filtered_bol, [selected_year]
            )
            scope31_supplier_specific_results = calculate_scope31_emissions_supplier_specific(
                spend_file, geo_content, filtered_awb, filtered_bol, [selected_year]
            )

    # ==================== EMISSIONS SUMMARY ====================
    st.header("Emissions Summary")

    # First calculate the basic AWB and BOL emissions (needed for TCE 3)
    awb_emissions = 0.0
    if filtered_awb is not None and not filtered_awb.empty and 'ghg_emissions_tCO2e' in filtered_awb.columns:
        awb_emissions = pd.to_numeric(
            filtered_awb.loc[filtered_awb['leg'] != 'TOTAL', 'ghg_emissions_tCO2e'],
            errors='coerce'
        ).sum()
        if pd.isna(awb_emissions):
            awb_emissions = 0.0

    bol_emissions = 0.0
    if filtered_bol is not None and not filtered_bol.empty and 'ghg_emissions_tCO2e' in filtered_bol.columns:
        bol_emissions = pd.to_numeric(filtered_bol['ghg_emissions_tCO2e'], errors='coerce').sum()
        if pd.isna(bol_emissions):
            bol_emissions = 0.0

    # Calculate TCE emissions
    tce1_results = calculate_tce1_breakdown(filtered_awb, filtered_bol, precomputed_mappings)
    tce2_results = calculate_tce2_breakdown(filtered_awb, filtered_bol)
    tce3_value = awb_emissions + bol_emissions
    tce4_results = calculate_tce4_breakdown(filtered_awb, filtered_bol)

    # ==================== ADD TCE COLUMNS TO AWB DATA ====================
    if filtered_awb is not None and not filtered_awb.empty:
        # Create a copy to avoid modifying the original
        awb_with_tce = filtered_awb.copy()
        
        # Add TCE columns - for AWB, we'll add these to the first leg of each shipment
        # Get first legs only (to avoid duplicating TCE 1,2,4 across multiple legs)
        first_leg_mask = (awb_with_tce['leg'] != 'TOTAL') & ~awb_with_tce.duplicated('awb_row', keep='first')
        total_mask = awb_with_tce['leg'] == 'TOTAL'
        
        # Initialize TCE columns
        awb_with_tce['TCE1_Emissions_tCO2e'] = 0.0
        awb_with_tce['TCE2_Emissions_tCO2e'] = 0.0  
        awb_with_tce['TCE4_Emissions_tCO2e'] = 0.0
        
        # Calculate TCE emissions for each first leg
        for idx, row in awb_with_tce[first_leg_mask].iterrows():
            if pd.notna(row['gross_weight']) and row['gross_weight'] > 0:
                # TCE 1: Inland transport to export hub
                continent = None
                if pd.notna(row['origin_lat']) and pd.notna(row['origin_lon']):
                    continent = get_continent_from_coords(row['origin_lat'], row['origin_lon'])
                elif pd.notna(row['origin']):
                    try:
                        origin_coords = get_airport_coords(row['origin'])
                        if None not in origin_coords:
                            continent = get_continent_from_coords(origin_coords[0], origin_coords[1])
                    except:
                        pass
                
                distance_km = get_distance_by_continent(continent) if continent else 0.0
                tce1_emissions = calculate_tce1_emissions(row['gross_weight'], distance_km)
                awb_with_tce.at[idx, 'TCE1_Emissions_tCO2e'] = tce1_emissions
                
                # TCE 2: Export hub operations
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                tce2_emissions = calculate_tce_emissions(row['gross_weight'], 'Air', cargo_type)
                awb_with_tce.at[idx, 'TCE2_Emissions_tCO2e'] = tce2_emissions
                
                # TCE 4: Import hub operations (same calculation as TCE 2 for simplicity)
                tce4_emissions = calculate_tce_emissions(row['gross_weight'], 'Air', cargo_type)
                awb_with_tce.at[idx, 'TCE4_Emissions_tCO2e'] = tce4_emissions
        
        # For TOTAL rows, sum up the TCE emissions from all legs
        for apo in awb_with_tce[total_mask]['APO'].unique():
            apo_mask = (awb_with_tce['APO'] == apo) & (awb_with_tce['leg'] != 'TOTAL')
            tce1_total = awb_with_tce.loc[apo_mask, 'TCE1_Emissions_tCO2e'].sum()
            tce2_total = awb_with_tce.loc[apo_mask, 'TCE2_Emissions_tCO2e'].sum()
            tce4_total = awb_with_tce.loc[apo_mask, 'TCE4_Emissions_tCO2e'].sum()
            
            total_row_mask = (awb_with_tce['APO'] == apo) & (awb_with_tce['leg'] == 'TOTAL')
            awb_with_tce.loc[total_row_mask, 'TCE1_Emissions_tCO2e'] = tce1_total
            awb_with_tce.loc[total_row_mask, 'TCE2_Emissions_tCO2e'] = tce2_total
            awb_with_tce.loc[total_row_mask, 'TCE4_Emissions_tCO2e'] = tce4_total
    else:
        # Handle case when there's no AWB data
        awb_with_tce = None

    # ==================== ADD TCE COLUMNS TO BOL DATA ====================
    if filtered_bol is not None and not filtered_bol.empty:
        # Create a copy to avoid modifying the original
        bol_with_tce = filtered_bol.copy()
        
        # Initialize TCE columns
        bol_with_tce['TCE1_Emissions_tCO2e'] = 0.0
        bol_with_tce['TCE2_Emissions_tCO2e'] = 0.0
        bol_with_tce['TCE4_Emissions_tCO2e'] = 0.0
        
        # Use precomputed city_to_continent mapping instead of loading geo data again
        city_to_continent = precomputed_mappings['city_to_continent']
        
        # Calculate TCE emissions for each BOL row
        for idx, row in bol_with_tce.iterrows():
            if pd.notna(row['Gross weight, kg']) and row['Gross weight, kg'] > 0:
                # TCE 1: Inland transport to export hub
                continent = None
                if 'Loading Port Continent' in row and pd.notna(row['Loading Port Continent']):
                    continent = row['Loading Port Continent']
                elif pd.notna(row['Port of loading']):
                    port_upper = row['Port of loading'].upper()
                    continent = city_to_continent.get(port_upper)
                
                distance_km = get_distance_by_continent(continent) if continent else 0.0
                tce1_emissions = calculate_tce1_emissions(row['Gross weight, kg'], distance_km)
                bol_with_tce.at[idx, 'TCE1_Emissions_tCO2e'] = tce1_emissions
                
                # TCE 2: Export hub operations
                cargo_type = 'Reefer' if row.get('Temperature Control') == 'YES' else 'Dry'
                tce2_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                bol_with_tce.at[idx, 'TCE2_Emissions_tCO2e'] = tce2_emissions
                
                # TCE 4: Import hub operations (same calculation as TCE 2 for simplicity)
                tce4_emissions = calculate_tce_emissions(row['Gross weight, kg'], 'Sea', cargo_type)
                bol_with_tce.at[idx, 'TCE4_Emissions_tCO2e'] = tce4_emissions

    # Scope 3.1 totals (with proper error handling)
    scope31_total = 0.0
    if scope31_results is not None and not scope31_results.empty and 'Supplier' in scope31_results.columns:
        total_row = scope31_results[scope31_results['Supplier'] == 'TOTAL']
        if not total_row.empty and 't CO2e' in total_row.columns:
            scope31_total = total_row['t CO2e'].values[0]

    scope31_exiobase_total = 0.0
    if scope31_exiobase_results is not None and not scope31_exiobase_results.empty and 'Supplier' in scope31_exiobase_results.columns:
        total_row = scope31_exiobase_results[scope31_exiobase_results['Supplier'] == 'TOTAL']
        if not total_row.empty and 'Exiobase_t_CO2e' in total_row.columns:
            scope31_exiobase_total = total_row['Exiobase_t_CO2e'].values[0]

    scope31_supplier_total = 0.0
    if scope31_supplier_specific_results is not None and not scope31_supplier_specific_results.empty and 'Supplier' in scope31_supplier_specific_results.columns:
        total_row = scope31_supplier_specific_results[scope31_supplier_specific_results['Supplier'] == 'TOTAL']
        if not total_row.empty and 't_CO2e' in total_row.columns:
            scope31_supplier_total = total_row['t_CO2e'].values[0]

    # Scope 3.1 section - Single header with reordered columns
    st.markdown("#### Scope 3.1 Emissions - Purchased Goods and Services")

    # Calculate intensities for each method
    epa_intensity = calculate_scope31_intensity(scope31_results, spend_file, [selected_year], "epa") if scope31_results is not None else None
    exiobase_intensity = calculate_scope31_intensity(scope31_exiobase_results, spend_file, [selected_year], "exiobase") if scope31_exiobase_results is not None else None
    supplier_intensity = calculate_scope31_intensity(scope31_supplier_specific_results, spend_file, [selected_year], "supplier") if scope31_supplier_specific_results is not None else None

    # Create columns in the new order: 1st Exiobase, 2nd Supplier-Specific, 3rd EPA
    scope31_col1, scope31_col2, scope31_col3 = st.columns(3)

    with scope31_col1:
        st.markdown("**Exiobase**")
        st.metric(
            "Country-Specific Factors", 
            f"{scope31_exiobase_total:,.2f}",
            help="A spend-based approach using country-specific emission factors from Exiobase."
        )
        st.caption("ton CO₂e", unsafe_allow_html=True)
        st.markdown("**Intensity**")
        intensity_text = f"{exiobase_intensity:,.1f} g CO₂e/$" if exiobase_intensity is not None else "N/A"
        st.markdown(intensity_text)

    with scope31_col2:
        st.markdown("**AKDN & Exiobase**")
        st.metric(
            "Supplier-Specific Factors", 
            f"{scope31_supplier_total:,.2f}",
            help="Combines AKDN supplier-specific emission factors with country-specific Exiobase fallback for most accurate estimation."
        )
        st.caption("ton CO₂e", unsafe_allow_html=True)
        st.markdown("**Intensity**")
        intensity_text = f"{supplier_intensity:,.1f} g CO₂e/$" if supplier_intensity is not None else "N/A"
        st.markdown(intensity_text)

    with scope31_col3:
        st.markdown("**EPA**")
        st.metric(
            "USA EPA Factors", 
            f"{scope31_total:,.2f}",
            help="A spend-based approach using USA EPA emission factors based on NAICS codes."
        )
        st.caption("ton CO₂e", unsafe_allow_html=True)
        st.markdown("**Intensity**")
        intensity_text = f"{epa_intensity:,.1f} g CO₂e/$" if epa_intensity is not None else "N/A"
        st.markdown(intensity_text)

    # Scope 3.4 section - TCE Breakdown
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("#### Scope 3.4 Emissions - Transportation Emissions")

    # Calculate Scope 3.4 intensity
    scope34_intensity = calculate_scope34_intensity(filtered_awb, filtered_bol)

    # Calculate subtotals for Air and Ocean emissions
    air_subtotal = (
        tce1_results.get('Air', 0.0) + 
        tce2_results.get('Air', 0.0) + 
        awb_emissions + 
        tce4_results.get('Air', 0.0)
    )

    ocean_subtotal = (
        tce1_results.get('Sea', 0.0) + 
        tce2_results.get('Sea', 0.0) + 
        bol_emissions + 
        tce4_results.get('Sea', 0.0)
    )

    scope34_grand_total = air_subtotal + ocean_subtotal

    # Calculate percentages for the speedometer
    air_percentage = (air_subtotal / scope34_grand_total * 100) if scope34_grand_total > 0 else 0
    ocean_percentage = (ocean_subtotal / scope34_grand_total * 100) if scope34_grand_total > 0 else 0

    # Create speedometer chart
    fig = go.Figure()

    # Add the speedometer gauge
    fig.add_trace(go.Indicator(
        mode="gauge",
        value=air_percentage,
        domain={'x': [0, 1], 'y': [0, 1]},
        title={'text': "Air vs Ocean Emissions", 'font': {'size': 24}},
        gauge={
            'axis': {'range': [0, 100]},
            'bar': {'color': "rgba(0,0,0,0)"},
            'steps': [
                {'range': [0, ocean_percentage], 'color': '#2D5889'},
                {'range': [ocean_percentage, 100], 'color': '#ED7136'}
            ],
            'bgcolor': "#f0f2f6"
        }
    ))

    fig.update_layout(
        height=400,
        margin=dict(l=50, r=50, t=100, b=50),
        font={'color': "darkblue", 'family': "Arial"},
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)'
    )

    # Grand Total above the chart (guard formatting when values are None)
    scope34_total_text = f"{scope34_grand_total:,.1f} ton CO₂e" if scope34_grand_total is not None else "N/A"
    scope34_intensity_text = f"{scope34_intensity:,.1f} g CO₂e/tkm" if scope34_intensity is not None else "N/A"

    st.markdown(f"""
    <div style='background-color: white; padding: 5px; border-radius: 10px; text-align: center; margin-bottom: 5px;'>
        <h2 style='margin: 0; color: #2c3e50;'>Scope 3.4 Total:  {scope34_total_text}</h2>
        <h2 style='margin: 0; color: #2c3e50;'>Intensity:  {scope34_intensity_text}</h2>
    </div>
    """, unsafe_allow_html=True)

    # Create columns for Ocean, Chart, and Air with bottom alignment
    ocean_col, chart_col, air_col = st.columns([1, 2, 1])
    
    with ocean_col:
        # Use container to help with bottom alignment
        ocean_container = st.container()
        with ocean_container:
            st.markdown(f"""
            <div style='text-align: center; padding: 15px; border-radius: 10px; background-color: white;'>
                <h1 style='margin: 0; color: white;'> </h1>
                <h2 style='margin: 0; color: #2D5889;'>Ocean</h2>
                <h2 style='margin: 0; color: #2D5889;'>{ocean_percentage:.1f}%</h2>
            </div>
            """, unsafe_allow_html=True)
        
        # Add spacer to push content to bottom
        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
    
    with chart_col:
        st.plotly_chart(fig, use_container_width=True)
    
    with air_col:
        # Use container to help with bottom alignment
        air_container = st.container()
        with air_container:
            st.markdown(f"""
            <div style='text-align: center; padding: 15px; border-radius: 10px; background-color: white;'>
                <h1 style='margin: 0; color: white;'> </h1>
                <h2 style='margin: 0; color: #ED7136;'>Air</h2>
                <h2 style='margin: 0; color: #ED7136;'>{air_percentage:.1f}%</h2>
            </div>
            """, unsafe_allow_html=True)
        
        # Add spacer to push content to bottom
        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)

    # Add some spacing
    #st.markdown("<br>", unsafe_allow_html=True)

    # Create a grid for TCE 1-4 with Air, Ocean, and Total as row headers
    st.markdown("**Transport Chain Elements (TCE) Breakdown**")
    st.caption("Unit: ton CO₂e")

    # Create a DataFrame for the TCE data
    tce_data = {
        'Mode': ['Air', 'Ocean', 'Total'],
        'TCE 1': [
            f"{tce1_results.get('Air', 0.0):,.2f}",
            f"{tce1_results.get('Sea', 0.0):,.2f}",
            f"{tce1_results.get('Total', 0.0):,.2f}"
        ],
        'TCE 2': [
            f"{tce2_results.get('Air', 0.0):,.2f}",
            f"{tce2_results.get('Sea', 0.0):,.2f}",
            f"{tce2_results.get('Total', 0.0):,.2f}"
        ],
        'TCE 3': [
            f"{awb_emissions:,.2f}",
            f"{bol_emissions:,.2f}",
            f"{tce3_value:,.2f}"
        ],
        'TCE 4': [
            f"{tce4_results.get('Air', 0.0):,.2f}",
            f"{tce4_results.get('Sea', 0.0):,.2f}",
            f"{tce4_results.get('Total', 0.0):,.2f}"
        ],
        'Subtotal': [
            # Air subtotal (sum of TCE1 Air + TCE2 Air + TCE3 Air + TCE4 Air)
            f"{tce1_results.get('Air', 0.0) + tce2_results.get('Air', 0.0) + awb_emissions + tce4_results.get('Air', 0.0):,.2f}",
            # Ocean subtotal (sum of TCE1 Sea + TCE2 Sea + TCE3 Sea + TCE4 Sea)
            f"{tce1_results.get('Sea', 0.0) + tce2_results.get('Sea', 0.0) + bol_emissions + tce4_results.get('Sea', 0.0):,.2f}",
            # Total subtotal (sum of all TCEs)
            f"{tce1_results.get('Total', 0.0) + tce2_results.get('Total', 0.0) + tce3_value + tce4_results.get('Total', 0.0):,.2f}"
        ]
    }

    tce_df = pd.DataFrame(tce_data)

    # Style the DataFrame for better presentation
    styled_tce = tce_df.style \
        .set_table_styles([
            {'selector': 'thead th', 'props': [('background-color', '#f0f2f6'), 
                                            ('font-weight', 'bold'),
                                            ('text-align', 'center')]},
            {'selector': 'tbody tr:nth-child(even)', 'props': [('background-color', '#f9f9f9')]},
            {'selector': 'tbody tr:nth-child(odd)', 'props': [('background-color', 'white')]},
            {'selector': 'td', 'props': [('text-align', 'right'), ('padding', '8px')]},
            {'selector': 'th', 'props': [('text-align', 'left'), ('padding', '8px')]},
            # Highlight the subtotal column
            {'selector': 'td:nth-child(6), th:nth-child(6)', 'props': [('font-weight', 'bold')]}
        ]) \
        .hide(axis='index')

    # Display the table
    st.table(styled_tce)

    # Add tooltips/descriptions below the table
    st.caption("""
    **TCE Definitions:**
    - **TCE 1**: Inland Transport (to export hub)
    - **TCE 2**: Export Hub Operations  
    - **TCE 3**: Main Transport (air/ocean)
    - **TCE 4**: Import Hub Operations
    """)

    st.markdown("---")    
    st.markdown("<br>", unsafe_allow_html=True)
    st.header("Detailed Scope 3.1 and Scope 3.4 Analysis")

    # ==================== SCOPE 3.1 SUMMARY AND COMPARISON ====================
    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"]:
        if scope31_results is not None or scope31_exiobase_results is not None or scope31_supplier_specific_results is not None:
            st.subheader("Scope 3.1")
            
            # Add general explanation
            with st.expander("ℹ️ About Scope 3.1 Calculation Methods"):
                st.markdown("""
                **Scope 3.1: Purchased Goods and Services**
                
                Three different calculation methods are available:
                
                **1. EPA Method** 
                - A spend-based approach using standardized emission factors from the US Environmental Protection Agency
                - Emission factors are based on the NAICS codes of PAHO's product categories
                - Not country-specific, assumes US average supply chains
                
                **2. Exiobase Method** 
                - A spend-based approach using the Exiobase multi-regional input-output database
                - Accounts for country-specific emission factors based on suppliers' locations
                - Considers the countries and sectors involved in PAHO's supply chain
                
                **3. Supplier-Specific + Exiobase** 
                - A spend-based hybrid approach that prioritizes supplier-specific emission factors when available
                - Supplier-specific emission factors from Aga Khan Development Network are used first
                - Exiobase emission factors are used as fallback when supplier-specific data is unavailable
                """)
            
            # Combined visualization for all three methods
            if (scope31_results is not None and scope31_exiobase_results is not None and
                scope31_supplier_specific_results is not None and 
                len(scope31_results) > 1 and len(scope31_exiobase_results) > 1 and 
                len(scope31_supplier_specific_results) > 1):
                
                # Prepare comparison data
                epa_total = scope31_results[scope31_results['Supplier'] == 'TOTAL']['t CO2e'].values[0]
                exiobase_total = scope31_exiobase_results[
                    scope31_exiobase_results['Supplier'] == 'TOTAL'
                ]['Exiobase_t_CO2e'].values[0]
                supplier_total = scope31_supplier_specific_results[
                    scope31_supplier_specific_results['Supplier'] == 'TOTAL'
                ]['t_CO2e'].values[0]
                
                # Create comparison data in the desired order: 1st Exiobase, 2nd Supplier-Specific, 3rd EPA
                comparison_df = pd.DataFrame({
                    'Method': ['Exiobase', 'Supplier-Specific + Exiobase', 'EPA'],
                    'Emissions (tCO₂e)': [exiobase_total, supplier_total, epa_total]
                })
                
                fig = px.bar(
                    comparison_df,
                    x='Method',
                    y='Emissions (tCO₂e)',
                    color_discrete_sequence=['#C0C0C0'],  # Silver color for all bars
                    title='Total Emissions Comparison: All Methods',
                    labels={'Emissions (tCO₂e)': 'Total Emissions (tCO₂e)'},
                    text_auto='.1f'  # This automatically adds data labels with 1 decimal place
                )
                
                # Customize the layout
                fig.update_layout(
                    showlegend=False,  # Remove legend since we have clear labels
                    xaxis_title=None,  # Remove x-axis title
                    yaxis_title='Total Emissions (tCO₂e)',
                    font=dict(size=12),
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)'
                )
                
                # Improve data labels appearance
                fig.update_traces(
                    textposition='outside',  # Place data labels above the bars
                    textfont=dict(size=12, color='black'),
                    opacity=0.8
                )
                
                # Adjust y-axis to accommodate the data labels
                max_emission = max(exiobase_total, supplier_total, epa_total)
                fig.update_yaxes(range=[0, max_emission * 1.1])  # Add 10% padding for labels
                
                st.plotly_chart(fig, use_container_width=True)

    # ==================== SCOPE 3.4 SECTION ====================
    st.subheader("Scope 3.4")
    show_combined_map(filtered_awb, filtered_bol, precomputed_mappings)

    # Create columns for the new visualizations
    col1, col2, col3 = st.columns(3)

    with col1:
        # ==================== TOP 5 FLIGHT ROUTES BY COUNT ====================
        st.subheader("Top 5 Routes")
        
        top_routes = get_top_flight_routes(filtered_awb, n=5)
        
        if not top_routes.empty:
            fig_routes = px.bar(
                top_routes,
                x='flight_count',
                y='route',
                orientation='h',
                color_discrete_sequence=['#C0C0C0'],  # Single color instead of color scale
                labels={
                    'flight_count': 'Number of Flights',
                    'route': ''  # Remove Y-axis label
                },
                hover_data=['ghg_emissions_tCO2e', 'avg_distance_km'],
                title=''
            )
            fig_routes.update_layout(
                yaxis={'categoryorder': 'total ascending', 'title': ''},  # Remove Y-axis title
                showlegend=False,
                height=400
            )
            st.plotly_chart(fig_routes, use_container_width=True)
            
            # Display data table
            with st.expander("View Route Details"):
                display_df = top_routes.copy()
                display_df['Flights'] = display_df['flight_count']
                display_df['Emissions (tCO₂e)'] = display_df['ghg_emissions_tCO2e'].round(2)
                display_df['Avg Distance (km)'] = display_df['avg_distance_km'].round(0)
                st.dataframe(display_df[['route', 'Flights', 'Emissions (tCO₂e)', 'Avg Distance (km)']], 
                            hide_index=True)
        else:
            st.info("No flight route data available")

    with col2:
        # ==================== TOP 5 SOURCE LOCATIONS BY COUNT ====================
        st.subheader("Top 5 Origins")
        
        # Only pass AWB data, not BOL data
        top_sources = get_top_source_locations(filtered_awb, n=5)
        
        if not top_sources.empty:
            # Create stacked bar chart for shipment counts
            fig_sources = px.bar(
                top_sources,
                x='shipment_count',
                y='location',
                orientation='h',
                title='',
                color_discrete_sequence=['#C0C0C0'],
                labels={
                    'shipment_count': 'Number of Shipments',
                    'location': ''  # Remove Y-axis label
                },
                hover_data=['ghg_emissions_tCO2e']
            )
            fig_sources.update_layout(
                yaxis={'categoryorder': 'total ascending', 'title': ''},  # Remove Y-axis title
                height=400,
                showlegend=True
            )
            st.plotly_chart(fig_sources, use_container_width=True)
            
            # Display data table
            with st.expander("View Source Details"):
                display_df = top_sources.copy()
                display_df['Total Shipments'] = display_df['shipment_count']
                display_df['Total Emissions (tCO₂e)'] = display_df['ghg_emissions_tCO2e'].round(2)
                
                # Calculate shipment composition percentages
                if 'Air' in display_df.columns and 'Ocean' in display_df.columns:
                    display_df['Air %'] = (display_df['Air'] / display_df['shipment_count'] * 100).round(1)
                    display_df['Ocean %'] = (display_df['Ocean'] / display_df['shipment_count'] * 100).round(1)
                    st.dataframe(display_df[['location', 'Total Shipments', 'Air', 'Ocean', 'Air %', 'Ocean %', 'Total Emissions (tCO₂e)']], 
                                hide_index=True)
                else:
                    st.dataframe(display_df[['location', 'Total Shipments', 'Total Emissions (tCO₂e)']], 
                                hide_index=True)
        else:
            st.info("No AWB origin data available")

    with col3:
        # ==================== TOP 5 AIRLINES BY COUNT ====================
        st.subheader("Top 5 Airlines")
        
        if filtered_awb is not None and not filtered_awb.empty and 'airline' in filtered_awb.columns:
            # Prepare airline data - count flights instead of summing emissions
            airline_counts = filtered_awb[filtered_awb['leg'] != 'TOTAL'].groupby('airline').agg({
                'leg': 'count',
                'ghg_emissions_tCO2e': 'sum',
                'distance_km': 'sum'
            }).rename(columns={
                'leg': 'flight_count',
                'distance_km': 'total_distance_km'
            }).reset_index()
            
            # Get top 5 airlines by flight count
            top_airlines = airline_counts.nlargest(5, 'flight_count')
            
            if not top_airlines.empty:
                # Create bar chart
                fig = px.bar(
                    top_airlines,
                    x='flight_count',
                    y='airline',
                    orientation='h',
                    color_discrete_sequence=['#C0C0C0'],  # Single color instead of color scale
                    labels={
                        'flight_count': 'Number of Flights',
                        'airline': ''  # Remove Y-axis label
                    },
                    hover_data=['ghg_emissions_tCO2e', 'total_distance_km'],
                    title=''
                )
                fig.update_layout(
                    yaxis={'categoryorder': 'total ascending', 'title': ''},  # Remove Y-axis title
                    showlegend=False,
                    height=400
                )
                st.plotly_chart(fig, use_container_width=True)
                
                # Display data table
                with st.expander("View Airline Details"):
                    display_df = top_airlines.copy()
                    display_df['Flights'] = display_df['flight_count']
                    display_df['Emissions (tCO₂e)'] = display_df['ghg_emissions_tCO2e'].round(2)
                    display_df['Total Distance (km)'] = display_df['total_distance_km'].round(0)
                    st.dataframe(display_df[['airline', 'Flights', 'Emissions (tCO₂e)', 'Total Distance (km)']], 
                                hide_index=True)
            else:
                st.info("No airline data available")
        else:
            st.info("No airline data available")

    # ==================== ADDITIONAL METRICS ROW ====================
    st.subheader("Breakdown by Transport Chain Element (TCE)")
    st.caption("Unit: ton CO₂e")

    # Calculate TCE 1 (inland transport emissions)
    tce1_results = calculate_tce1_breakdown(filtered_awb, filtered_bol, precomputed_mappings)
    tce1_value = tce1_results['Total']

    # Calculate TCE 2 (export hub operations)
    tce2_results = calculate_tce2_breakdown(filtered_awb, filtered_bol)
    tce2_value = tce2_results['Total']

    # Calculate TCE 3 (same as sum of Air Transport Emissions and Ocean Transport Emissions)
    tce_3_value = awb_emissions + bol_emissions

    # Calculate TCE 4 (import hub operations)
    tce4_results = calculate_tce4_breakdown(filtered_awb, filtered_bol)
    tce4_value = tce4_results['Total']

    # Create metrics columns
    metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)

    with metric_col1:
        st.metric("TCE 1", f"{tce1_value:,.1f}")

    with metric_col2:
        st.metric("TCE 2", f"{tce2_value:,.1f}")

    with metric_col3:
        st.metric("TCE 3", f"{tce_3_value:,.1f}")

    with metric_col4:
        st.metric("TCE 4", f"{tce4_value:,.1f}")

    # ==================== TABS DISPLAY ====================
    # Create tabs for all data tables
    tab_list = []
    
    if has_awb:
        tab_list.append("Air Waybill (AWB)")
    if has_bol:
        tab_list.append("Bill of Lading (BOL)")
    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"]:
        tab_list.extend(["EPA Method", "Exiobase Method", "Supplier-Specific + Exiobase"])
    
    if tab_list:
        tabs = st.tabs(tab_list)
        tab_index = 0
    else:
        tabs = []
        tab_index = 0

    # ==================== MODIFIED AWB TAB ====================
    if has_awb:
        with tabs[tab_index]:
            st.header("Air Waybill (AWB) Analysis")
            
            if awb_with_tce is not None and not awb_with_tce.empty:
                st.success(f"✅ Processed {len(awb_with_tce)} AWB records")
                
                st.subheader("Flight Data with TCE Breakdown")
                
                # Reorder columns to show TCE emissions at the end
                column_order = [
                    col for col in awb_with_tce.columns 
                    if col not in ['TCE1_Emissions_tCO2e', 'TCE2_Emissions_tCO2e', 'TCE4_Emissions_tCO2e']
                ] + ['TCE1_Emissions_tCO2e', 'TCE2_Emissions_tCO2e', 'TCE4_Emissions_tCO2e']
                
                styled_awb = style_awb_dataframe(awb_with_tce[column_order])
                st.dataframe(styled_awb)
                
                # Update download function to include TCE columns
                create_excel_download(awb_with_tce[column_order], "awb_emissions_with_tce.xlsx")

            else:
                st.info("No valid AWB data found" + (" (filtered out)" if selected_suppliers or selected_apos else ""))
        tab_index += 1

    # ==================== MODIFIED BOL TAB ====================
    if has_bol:
        with tabs[tab_index]:
            st.header("Bill of Lading (BOL) Analysis")
            
            if bol_with_tce is not None and not bol_with_tce.empty:
                st.success(f"✅ Processed {len(bol_with_tce)} BOL records")
                
                st.subheader("Shipping Data with TCE Breakdown")
                
                # Reorder columns to show TCE emissions at the end
                column_order = [
                    col for col in bol_with_tce.columns 
                    if col not in ['TCE1_Emissions_tCO2e', 'TCE2_Emissions_tCO2e', 'TCE4_Emissions_tCO2e']
                ] + ['TCE1_Emissions_tCO2e', 'TCE2_Emissions_tCO2e', 'TCE4_Emissions_tCO2e']
                
                st.dataframe(bol_with_tce[column_order])
                
                # Update download function to include TCE columns
                create_excel_download(bol_with_tce[column_order], "bol_emissions_with_tce.xlsx")
                
            else:
                st.error("❌ No valid BOL data found" + (" (filtered out)" if selected_suppliers or selected_apos else ""))
        tab_index += 1

    # ==================== SCOPE 3.1 DETAILED TABLES ====================
    if analysis_scope in ["Scope 3.1 (Purchased Goods)", "Both"]:
        # EPA Method Tab
        with tabs[tab_index]:
            st.header("EPA Method Details")
            
            if scope31_results is not None:
                formatted_df = scope31_results.copy()
                if 'Supplier' in formatted_df.columns:
                    formatted_df['Supplier'] = formatted_df['Supplier'].fillna('Unknown')
                
                st.dataframe(
                    formatted_df.style.format({
                        'Amount (USD)': '{:,.2f}',
                        'kg CO2e': '{:,.2f}',
                        't CO2e': '{:,.6f}',
                        'kg CO2e/USD': '{:,.6f}'
                    })
                )
                
                # Download button for EPA results
                create_excel_download(scope31_results, "scope31_epa_results.xlsx")
            else:
                st.info("No EPA results available")
        tab_index += 1
        
        # Exiobase Method Tab
        with tabs[tab_index]:
            st.header("Exiobase Method Details")
            
            if scope31_exiobase_results is not None:
                formatted_df = scope31_exiobase_results.copy()
                if 'Supplier' in formatted_df.columns:
                    formatted_df['Supplier'] = formatted_df['Supplier'].fillna('Unknown')
                
                st.dataframe(
                    formatted_df.style.format({
                        'Amount (USD)': '{:,.2f}',
                        'Exiobase_kg_CO2e': '{:,.2f}',
                        'Exiobase_t_CO2e': '{:,.6f}',
                        'Exiobase_kg_CO2e/USD': '{:,.6f}'
                    })
                )
                
                # Download button for Exiobase results
                create_excel_download(scope31_exiobase_results, "scope31_exiobase_results.xlsx")
            else:
                st.info("No Exiobase results available")
        tab_index += 1
        
        # Supplier-Specific + Exiobase Tab
        with tabs[tab_index]:
            st.header("Supplier-Specific + Exiobase Details")
            
            if scope31_supplier_specific_results is not None:
                formatted_df = scope31_supplier_specific_results.copy()
                if 'Supplier' in formatted_df.columns:
                    formatted_df['Supplier'] = formatted_df['Supplier'].fillna('Unknown')
                
                st.dataframe(formatted_df)
                
                # Download button for Supplier-Specific results
                create_excel_download(scope31_supplier_specific_results, "scope31_supplier_specific_results.xlsx")
            else:
                st.info("No Supplier-Specific results available")

    # ==================== REPORT SECTION ====================
    # Aggregate report metrics for the bottom section
    report_year = selected_year
    report_scope31_total = scope31_supplier_total
    report_scope34_total = scope34_grand_total

    # Scope 3.1 activity volume and unique purchase orders
    report_spend_amount = None
    report_po_count = 0
    if spend_file is not None:
        try:
            spend_df = pd.read_excel(
                spend_file,
                usecols=["YearReceipt", "FundType", "LineDescription", "Amount per PO Line", "PurchaseOrderNumber"]
            )
            report_spend_filtered = spend_df.copy()
            if report_year:
                report_spend_filtered = report_spend_filtered[
                    report_spend_filtered['YearReceipt'].astype(str).isin([report_year])
                ]
            report_spend_filtered = report_spend_filtered[
                report_spend_filtered['FundType'].isin(['Revolving Fund', 'Strategic Fund'])
            ]
            report_spend_filtered = report_spend_filtered[
                ~report_spend_filtered['LineDescription'].str.contains('donations|discount', case=False, na=False)
            ]
            report_spend_amount = report_spend_filtered['Amount per PO Line'].sum()
            report_po_count = report_spend_filtered['PurchaseOrderNumber'].nunique()
        except Exception:
            report_spend_amount = None
            report_po_count = 0

    # Scope 3.4 activity volume and shipment count
    scope34_weight_ton = 0.0
    scope34_ton_km = 0.0
    awb_shipments = 0
    bol_shipments = 0

    if filtered_awb is not None and not filtered_awb.empty:
        awb_df = filtered_awb.copy()
        if 'gross_weight_ton' in awb_df.columns:
            awb_df['gross_weight_ton'] = pd.to_numeric(awb_df['gross_weight_ton'], errors='coerce')
        else:
            awb_df['gross_weight_ton'] = pd.to_numeric(awb_df.get('gross_weight', None), errors='coerce') / 1000
        awb_df['distance_km'] = pd.to_numeric(awb_df.get('distance_km', None), errors='coerce')
        awb_df = awb_df[(awb_df['gross_weight_ton'] > 0) & (awb_df['distance_km'] > 0)]
        scope34_weight_ton += awb_df['gross_weight_ton'].sum()
        scope34_ton_km += (awb_df['gross_weight_ton'] * awb_df['distance_km']).sum()
        if 'awb_row' in filtered_awb.columns:
            awb_shipments = filtered_awb['awb_row'].nunique()
        else:
            awb_shipments = len(filtered_awb)

    if filtered_bol is not None and not filtered_bol.empty:
        bol_df = filtered_bol.copy()
        bol_df['Gross weight, kg'] = pd.to_numeric(bol_df.get('Gross weight, kg', None), errors='coerce')
        bol_df['Sea Distance (km)'] = pd.to_numeric(bol_df.get('Sea Distance (km)', None), errors='coerce')
        bol_valid = bol_df[(bol_df['Gross weight, kg'] > 0) & (bol_df['Sea Distance (km)'] > 0)]
        scope34_weight_ton += (bol_valid['Gross weight, kg'] / 1000).sum()
        scope34_ton_km += ((bol_valid['Gross weight, kg'] / 1000) * bol_valid['Sea Distance (km)']).sum()
        bol_shipments = 0
        for col in ['Bill of Lading No.', 'Bill of Lading Number', 'BOL No.', 'BOL Number', 'Bill of Lading', 'BOL']:
            if col in filtered_bol.columns:
                bol_shipments = filtered_bol[col].nunique()
                break
        if bol_shipments == 0:
            bol_shipments = len(filtered_bol)

    report_scope34_activity = None
    if scope34_weight_ton > 0 or scope34_ton_km > 0:
        report_scope34_activity = f"{scope34_ton_km:,.0f} t·km + {scope34_weight_ton:,.0f} t"

    report_scope31_activity = None
    if report_spend_amount is not None:
        report_scope31_activity = f"{report_spend_amount:,.2f} USD"

    report_scope31_records = f"{report_po_count} purchase order{'s' if report_po_count != 1 else ''}"
    report_scope34_records = f"{awb_shipments + bol_shipments} shipments"

    # Compute TCE1 fixed-distance ton-km so transport intensity includes TCE1 + TCE3
    _tce1_ton_km = 0.0
    if filtered_awb is not None and not filtered_awb.empty:
        _awb_ref = awb_with_tce if awb_with_tce is not None else filtered_awb
        _mask = (_awb_ref['leg'] != 'TOTAL') & ~_awb_ref.duplicated('awb_row', keep='first')
        for _, _r in _awb_ref[_mask].iterrows():
            _wt = pd.to_numeric(_r.get('gross_weight', 0), errors='coerce') / 1000
            if _wt > 0:
                _cont = None
                if pd.notna(_r.get('origin_lat')) and pd.notna(_r.get('origin_lon')):
                    _cont = get_continent_from_coords(_r['origin_lat'], _r['origin_lon'])
                _tce1_ton_km += _wt * get_distance_by_continent(_cont)
    if filtered_bol is not None and not filtered_bol.empty:
        _bol_ref = bol_with_tce if bol_with_tce is not None else filtered_bol
        for _, _r in _bol_ref.iterrows():
            _wt = pd.to_numeric(_r.get('Gross weight, kg', 0), errors='coerce') / 1000
            if _wt > 0:
                _cont = _r.get('Loading Port Continent')
                if pd.isna(_cont) and pd.notna(_r.get('Port of loading')):
                    _cont = precomputed_mappings['city_to_continent'].get(str(_r['Port of loading']).upper())
                _tce1_ton_km += _wt * get_distance_by_continent(_cont)
    _headline_transport_ton_km = scope34_ton_km + _tce1_ton_km

    transport_intensity_component1 = None
    transport_intensity_component2 = None
    if _headline_transport_ton_km > 0:
        transport_intensity_component1 = (
            (tce1_results.get('Total', 0.0) + tce_3_value) * 1_000_000
        ) / _headline_transport_ton_km
    if scope34_weight_ton > 0:
        transport_intensity_component2 = ((tce2_results.get('Total', 0.0) + tce4_results.get('Total', 0.0)) * 1_000_000) / (scope34_weight_ton*2)

    report_scope31_intensity = None
    if supplier_intensity is not None:
        report_scope31_intensity = f"{supplier_intensity:,.1f} g CO₂e/$"
    report_scope34_intensity = None
    if transport_intensity_component1 is not None and transport_intensity_component2 is not None:
        report_scope34_intensity = (
            f"{transport_intensity_component1:,.1f} g CO₂e/t-km (transport); "
            f"{transport_intensity_component2:,.1f} g CO₂e/t (hubs)"
        )

    report_df = pd.DataFrame({
        'Indicator': [
            'Total GHG emissions (t CO₂e)',
            'Activity volume',
            'Number of activity records',
            'Emission intensity',
            'Methodology',
            'Base year'
        ],
        'Scope 3 Cat. 1': [
            f"{report_scope31_total:,.2f}",
            report_scope31_activity or 'N/A',
            report_scope31_records,
            report_scope31_intensity or 'N/A',
            'Spend-based (3 levels): USA EPA, EXIOBASE, AKDN supplier-specific',
            '2023'
        ],
        'Scope 3 Cat. 4': [
            f"{report_scope34_total:,.2f}",
            report_scope34_activity or 'N/A',
            report_scope34_records,
            report_scope34_intensity or 'N/A',
            'Distance-based per ISO 14083 (TCE 1–4); GLEC 2025 factors',
            'To confirm — 2023 or 2024'
        ]
    })

    st.markdown("---")
    with st.expander("📊 Report-year headline results", expanded=False):
        st.table(report_df)

        combined_total = report_scope31_total + report_scope34_total
        combined_percentage_3 = 0.0
        combined_percentage_4 = 0.0
        if combined_total > 0:
            combined_percentage_3 = (report_scope31_total / combined_total) * 100
            combined_percentage_4 = (report_scope34_total / combined_total) * 100

        st.markdown(
            f"Combined Scope 3 (Categories 1 + 4) emissions for {report_year} are estimated at **{combined_total:,.2f}**, "
            f"with Category 1 contributing approximately **{combined_percentage_3:,.1f}%** and Category 4 approximately "
            f"**{combined_percentage_4:,.1f}%** of the consolidated upstream supply-chain footprint."
        )

    # ==================== SECOND REPORT SECTION ====================
    with st.expander("📈 Scope 3 Category 1 – Results by methodology level", expanded=False):
        # ====== Methodology Comparison Table ======
        st.subheader("Results by Methodology Level")
        
        # Extract totals and spend for each methodology
        methodology_data = {
            'Level': ['Level 1', 'Level 2', 'Level 3', 'Reported (Level 3 – preferred)'],
            'Emission factor source': [
                'US EPA Supply Chain v1.3 (NAICS-6)',
                'EXIOBASE v3.9.4 (country-specific)',
                'AKDN supplier-specific + EXIOBASE',
                'Hybrid: AKDN + EXIOBASE'
            ],
            'Total spend (US$)': [
                f"{report_spend_amount:,.2f}" if report_spend_amount is not None else 'N/A',
                f"{report_spend_amount:,.2f}" if report_spend_amount is not None else 'N/A',
                f"{report_spend_amount:,.2f}" if report_spend_amount is not None else 'N/A',
                f"{report_spend_amount:,.2f}" if report_spend_amount is not None else 'N/A'
            ],
            'GHG emissions (t CO₂e)': [
                f"{scope31_total:,.2f}",
                f"{scope31_exiobase_total:,.2f}",
                f"{scope31_supplier_total:,.2f}",
                f"{scope31_supplier_total:,.2f}"
            ]
        }
        methodology_df = pd.DataFrame(methodology_data)
        st.table(methodology_df)
        
        # ====== Text Summary with Top Category ======
        if scope31_supplier_specific_results is not None and not scope31_supplier_specific_results.empty:
            # Find top spending category by emissions for Level 3 (supplier-specific)
            supplier_df = scope31_supplier_specific_results[scope31_supplier_specific_results['Supplier'] != 'TOTAL'].copy()
            if 'Spend Category' in supplier_df.columns:
                category_emissions = supplier_df.groupby('Spend Category')['kg_CO2e'].sum().sort_values(ascending=False)
            else:
                category_emissions = supplier_df.groupby('SpendCategory')['kg_CO2e'].sum().sort_values(ascending=False)
            
            if len(category_emissions) > 0:
                top_category = category_emissions.index[0]
                top_category_emissions = category_emissions.iloc[0] / 1000  # convert to t CO2e
                total_emissions_g = supplier_df['kg_CO2e'].sum()
                top_category_pct = (category_emissions.iloc[0] / total_emissions_g * 100) if total_emissions_g > 0 else 0
                
                st.markdown(
                    f"The Level 3 hybrid approach is the value reported for Category 1. Across all three levels, the "
                    f"**{top_category}** category accounts for the largest share of emissions "
                    f"(**{top_category_pct:.1f}%**)."
                )
        
        # ====== Top 10 Product Categories (Level 3) ======
        st.subheader("Top Product Categories (Level 3)")
        
        if scope31_supplier_specific_results is not None and not scope31_supplier_specific_results.empty:
            category_df = scope31_supplier_specific_results[scope31_supplier_specific_results['Supplier'] != 'TOTAL'].copy()
            if 'Spend Category' in category_df.columns:
                cat_col = 'Spend Category'
            else:
                cat_col = 'SpendCategory'
            
            category_totals = category_df.groupby(cat_col)['kg_CO2e'].sum().sort_values(ascending=False)
            category_totals_t = category_totals / 1000  # convert to t CO2e
            
            total_cat_emissions = category_totals_t.sum()
            
            # Get top 10 categories
            top_10_categories = category_totals_t.head(10)
            other_total = category_totals_t.iloc[10:].sum() if len(category_totals_t) > 10 else 0.0
            
            # Build the table
            category_table_data = []
            for idx, (cat_name, emissions) in enumerate(top_10_categories.items(), 1):
                pct = (emissions / total_cat_emissions * 100) if total_cat_emissions > 0 else 0
                category_table_data.append({
                    'PAHO Product Category': cat_name,
                    'GHG (t CO₂e)': f"{emissions:,.2f}",
                    '% of total': f"{pct:.1f}%"
                })
            
            # Add "Other categories" row
            if other_total > 0:
                other_pct = (other_total / total_cat_emissions * 100) if total_cat_emissions > 0 else 0
                category_table_data.append({
                    'PAHO Product Category': 'Other categories (combined)',
                    'GHG (t CO₂e)': f"{other_total:,.2f}",
                    '% of total': f"{other_pct:.1f}%"
                })
            
            # Add Total row
            category_table_data.append({
                'PAHO Product Category': 'Total',
                'GHG (t CO₂e)': f"{total_cat_emissions:,.2f}",
                '% of total': '100%'
            })
            
            category_table_df = pd.DataFrame(category_table_data)
            st.table(category_table_df)
        
        # ====== Top Suppliers (Level 3) ======
        st.subheader("Top Suppliers (Level 3)")
        
        if scope31_supplier_specific_results is not None and not scope31_supplier_specific_results.empty:
            supplier_df = scope31_supplier_specific_results[scope31_supplier_specific_results['Supplier'] != 'TOTAL'].copy()
            
            supplier_totals = supplier_df.groupby('Supplier')['kg_CO2e'].sum().sort_values(ascending=False)
            supplier_totals_t = supplier_totals / 1000  # convert to t CO2e
            
            total_supplier_emissions = supplier_totals_t.sum()
            
            # Get top 11 suppliers
            top_11_suppliers = supplier_totals_t.head(11)
            other_total = supplier_totals_t.iloc[11:].sum() if len(supplier_totals_t) > 11 else 0.0
            
            # Build the table
            supplier_table_data = []
            for idx, (supplier_name, emissions) in enumerate(top_11_suppliers.items(), 1):
                pct = (emissions / total_supplier_emissions * 100) if total_supplier_emissions > 0 else 0
                supplier_table_data.append({
                    'Supplier': supplier_name,
                    'GHG (t CO₂e)': f"{emissions:,.2f}",
                    '% of total': f"{pct:.1f}%"
                })
            
            # Add "Other suppliers" row
            if other_total > 0:
                other_pct = (other_total / total_supplier_emissions * 100) if total_supplier_emissions > 0 else 0
                supplier_table_data.append({
                    'Supplier': 'Other suppliers (combined)',
                    'GHG (t CO₂e)': f"{other_total:,.2f}",
                    '% of total': f"{other_pct:.1f}%"
                })
            
            # Add Total row
            supplier_table_data.append({
                'Supplier': 'Total',
                'GHG (t CO₂e)': f"{total_supplier_emissions:,.2f}",
                '% of total': '100%'
            })
            
            supplier_table_df = pd.DataFrame(supplier_table_data)
            st.table(supplier_table_df)

    # ==================== THIRD REPORT SECTION ====================
    with st.expander("📊 Scope 3 Category 4 — Results by Transport Chain Element", expanded=False):
        tce1_total = tce1_results.get('Total', 0.0)
        tce2_total = tce2_results.get('Total', 0.0)
        tce3_total = tce3_value
        tce4_total = tce4_results.get('Total', 0.0)
        scope34_total = scope34_grand_total

        air_tce3 = awb_emissions
        sea_tce3 = bol_emissions
        tce1_modes = tce1_results.get('Total', 0.0)
        hub_operations = tce2_total + tce4_total

        awb_weight_ton = 0.0
        awb_ton_km = 0.0
        if filtered_awb is not None and not filtered_awb.empty:
            awb_df = filtered_awb[filtered_awb['leg'] != 'TOTAL'].copy()
            if 'gross_weight_ton' in awb_df.columns:
                awb_df['gross_weight_ton'] = pd.to_numeric(awb_df['gross_weight_ton'], errors='coerce')
            else:
                awb_df['gross_weight_ton'] = pd.to_numeric(awb_df.get('gross_weight', None), errors='coerce') / 1000
            awb_df['distance_km'] = pd.to_numeric(awb_df.get('distance_km', None), errors='coerce')
            awb_df = awb_df[(awb_df['gross_weight_ton'] > 0) & (awb_df['distance_km'] > 0)]
            awb_weight_ton = awb_df['gross_weight_ton'].sum()
            awb_ton_km = (awb_df['gross_weight_ton'] * awb_df['distance_km']).sum()

        bol_weight_ton = 0.0
        bol_ton_km = 0.0
        if filtered_bol is not None and not filtered_bol.empty:
            bol_df = filtered_bol.copy()
            bol_df['Gross weight, kg'] = pd.to_numeric(bol_df.get('Gross weight, kg', None), errors='coerce')
            bol_df['Sea Distance (km)'] = pd.to_numeric(bol_df.get('Sea Distance (km)', None), errors='coerce')
            bol_valid = bol_df[(bol_df['Gross weight, kg'] > 0) & (bol_df['Sea Distance (km)'] > 0)]
            bol_weight_ton = (bol_valid['Gross weight, kg'] / 1000).sum()
            bol_ton_km = ((bol_valid['Gross weight, kg'] / 1000) * bol_valid['Sea Distance (km)']).sum()

        fixed_tce1_ton_km = 0.0
        if filtered_awb is not None and not filtered_awb.empty:
            awb_ref = awb_with_tce if awb_with_tce is not None else filtered_awb
            first_leg_mask = (awb_ref['leg'] != 'TOTAL') & ~awb_ref.duplicated('awb_row', keep='first')
            for _, row in awb_ref[first_leg_mask].iterrows():
                weight_ton = 0.0
                if pd.notna(row.get('gross_weight', None)):
                    weight_ton = pd.to_numeric(row.get('gross_weight', 0), errors='coerce') / 1000
                if weight_ton <= 0:
                    continue
                continent = None
                if pd.notna(row.get('origin_lat')) and pd.notna(row.get('origin_lon')):
                    continent = get_continent_from_coords(row['origin_lat'], row['origin_lon'])
                elif pd.notna(row.get('origin')):
                    try:
                        origin_coords = get_airport_coords(row['origin'])
                        if None not in origin_coords:
                            continent = get_continent_from_coords(origin_coords[0], origin_coords[1])
                    except Exception:
                        pass
                fixed_tce1_ton_km += weight_ton * get_distance_by_continent(continent)

        if filtered_bol is not None and not filtered_bol.empty:
            bol_ref = bol_with_tce if bol_with_tce is not None else filtered_bol
            for _, row in bol_ref.iterrows():
                weight_ton = 0.0
                if pd.notna(row.get('Gross weight, kg', None)):
                    weight_ton = pd.to_numeric(row.get('Gross weight, kg', 0), errors='coerce') / 1000
                if weight_ton <= 0:
                    continue
                continent = None
                if pd.notna(row.get('Loading Port Continent')):
                    continent = row['Loading Port Continent']
                elif pd.notna(row.get('Port of loading')):
                    port_upper = str(row['Port of loading']).upper()
                    continent = precomputed_mappings['city_to_continent'].get(port_upper)
                fixed_tce1_ton_km += weight_ton * get_distance_by_continent(continent)

        total_transport_activity_ton_km = fixed_tce1_ton_km + awb_ton_km + bol_ton_km
        tce1_weight_distance = fixed_tce1_ton_km
        total_shipment_weight_ton = awb_weight_ton + bol_weight_ton

        overall_intensity = None
        transport_intensity = None
        air_intensity = None
        sea_intensity = None
        mixed_mode_tce1_intensity = None
        hubs_intensity = None

        if total_transport_activity_ton_km > 0:
            overall_intensity = (scope34_total * 1_000_000) / total_transport_activity_ton_km
            transport_intensity = ((tce1_total + tce3_total) * 1_000_000) / total_transport_activity_ton_km
        if awb_ton_km > 0:
            air_intensity = (air_tce3 * 1_000_000) / awb_ton_km
        if bol_ton_km > 0:
            sea_intensity = (sea_tce3 * 1_000_000) / bol_ton_km
        if tce1_weight_distance > 0:
            mixed_mode_tce1_intensity = (tce1_total * 1_000_000) / tce1_weight_distance
        if total_shipment_weight_ton > 0:
            hubs_intensity = ((tce2_total + tce4_total) * 1_000_000) / (total_shipment_weight_ton*2)

        tce_table = pd.DataFrame({
            'TCE': ['TCE 1', 'TCE 2', 'TCE 3', 'TCE 4', 'Total'],
            'Segment': [
                'Supplier → Port of Origin (mixed modes)',
                'Hub operations at Port of Origin',
                'Port of Origin → Port of Destination (air + sea)',
                'Hub operations at Port of Destination',
                'All transport chain elements (TCE 1–4)'
            ],
            'GHG (t CO₂e)': [
                f"{tce1_total:,.2f}",
                f"{tce2_total:,.2f}",
                f"{tce3_total:,.2f}",
                f"{tce4_total:,.2f}",
                f"{scope34_total:,.2f}"
            ],
            '% of total': [
                f"{(tce1_total / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(tce2_total / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(tce3_total / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(tce4_total / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                '100%'
            ]
        })

        mode_table = pd.DataFrame({
            'Mode': ['Air (TCE 3)', 'Sea (TCE 3)', 'Combinations of modes (TCE 1)', 'Hub operations (TCE 2 + TCE 4)', 'Total'],
            'GHG (t CO₂e)': [
                f"{air_tce3:,.2f}",
                f"{sea_tce3:,.2f}",
                f"{tce1_modes:,.2f}",
                f"{hub_operations:,.2f}",
                f"{scope34_total:,.2f}"
            ],
            '% of transport total': [
                f"{(air_tce3 / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(sea_tce3 / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(tce1_modes / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                f"{(hub_operations / scope34_total * 100):.1f}%" if scope34_total > 0 else 'N/A',
                '100%'
            ]
        })

        intensity_table = pd.DataFrame({
            'Intensity indicator': [
                'Overall transport + hubs intensity (avg.)',
                'Transport intensity (avg.)',
                'Air mode (TCE 3, Great Circle Distance)',
                'Sea mode (TCE 3, Shortest Feasible Distance)',
                'Mixed-mode TCE 1',
                'Hubs intensity (TCE 2 + TCE 4)'
            ],
            'Value': [
                f"{overall_intensity:,.1f} g CO₂e/t-km" if overall_intensity is not None else 'N/A',
                f"{transport_intensity:,.1f} g CO₂e/t-km" if transport_intensity is not None else 'N/A',
                f"{air_intensity:,.1f} g CO₂e/t-km" if air_intensity is not None else 'N/A',
                f"{sea_intensity:,.1f} g CO₂e/t-km" if sea_intensity is not None else 'N/A',
                f"{mixed_mode_tce1_intensity:,.1f} g CO₂e/t-km" if mixed_mode_tce1_intensity is not None else 'N/A',
                f"{hubs_intensity:,.1f} g CO₂e/t" if hubs_intensity is not None else 'N/A'
            ]
        })

        st.subheader("Scope 3 Category 4 — Results by Transport Chain Element")
        st.table(tce_table)

        st.subheader("Category 4 — Disaggregation by transport mode")
        st.table(mode_table)

        st.subheader("Category 4 — Emission intensities")
        st.table(intensity_table)

        total_transport_text = f"{total_transport_activity_ton_km:,.0f} t·km"
        total_hub_throughput = f"{total_shipment_weight_ton:,.0f} t"
        st.markdown(
            f"Total transport activity for the reporting year: **{total_transport_text}** (TCE 1 + TCE 3); "
            f"total hub throughput: **{total_hub_throughput}** (TCE 2 + TCE 4). — *Note: distance for TCE 1 = fixed distances by continent (e.g., 662km for Asia, 100km for Europe, etc.); distance for TCE 3 = calculated distance from AWB/BOL origin/destination data.*"
        )

    # ==================== FOURTH REPORT SECTION ====================
    kpi_cat1_intensity = None
    if report_spend_amount is not None and report_spend_amount > 0:
        kpi_cat1_intensity = report_scope31_total / report_spend_amount

    kpi_cat4_transport_intensity = None
    if scope34_ton_km > 0:
        kpi_cat4_transport_intensity = ((tce1_results.get('Total', 0.0) + tce3_value) * 1_000_000) / scope34_ton_km

    kpi_cat4_hubs_intensity = None
    if scope34_weight_ton > 0:
        kpi_cat4_hubs_intensity = ((tce2_results.get('Total', 0.0) + tce4_results.get('Total', 0.0)) * 1_000_000) / scope34_weight_ton

    consolidated_df = pd.DataFrame({
        'Category': [
            'Scope 3 Cat. 1 — Purchased Goods & Services',
            'Scope 3 Cat. 4 — Upstream Transportation & Distribution',
            'TOTAL — Reported Scope 3 (Cat. 1 + Cat. 4)'
        ],
        'GHG emissions (t CO₂e)': [
            f"{report_scope31_total:,.2f}",
            f"{report_scope34_total:,.2f}",
            f"{combined_total:,.2f}"
        ]
    })

    if kpi_cat1_intensity is not None:
        kpi2_value = f"{report_scope31_total:,.2f} t CO₂e / US${report_spend_amount:,.2f} = {kpi_cat1_intensity:,.6f} t CO₂e / US$"
    else:
        kpi2_value = 'N/A'

    if kpi_cat4_transport_intensity is not None and kpi_cat4_hubs_intensity is not None:
        kpi3_value = (
            f"Transport: {kpi_cat4_transport_intensity:,.1f} g CO₂e/t-km; "
            f"Hubs: {kpi_cat4_hubs_intensity:,.1f} g CO₂e/t"
        )
    else:
        kpi3_value = 'N/A'

    kpi_df = pd.DataFrame({
        'KPI': [
            'KPI 1 — Absolute emissions',
            'KPI 2 — Economic intensity (Cat. 1)',
            'KPI 3 — Physical intensity (Cat. 4)',
            'KPI 4 — Physical intensity (Cat. 1)'
        ],
        'Value': [
            f"Cat. 1: {report_scope31_total:,.2f} t CO₂e; Cat. 4: {report_scope34_total:,.2f} t CO₂e; Combined: {combined_total:,.2f} t CO₂e",
            kpi2_value,
            kpi3_value,
            'Σ Emissions / Vaccine doses (t CO₂e / dose) — to be added once cradle-to-gate factors become available from suppliers.'
        ]
    })

    with st.expander("📊 Consolidated Scope 3 (Categories 1 + 4)", expanded=False):
        st.subheader("Consolidated Scope 3 (Categories 1 + 4)")
        st.table(consolidated_df)
        st.subheader("Performance Indicators (KPIs)")
        st.table(kpi_df)

    # ==================== WORD DOCUMENT DOWNLOAD ====================
    st.markdown("---")
    
    def generate_word_report():
        """Generate populated Word document with all calculated metrics and top-N results."""
        try:
            # Load template
            template_path = os.path.join(BASE_DIR, 'PAHO GHG Accounting Report 2025 - Scope 3 Cat 1 & Cat 4_TEMPLATE_PLACEHOLDERS.docx')
            if not os.path.exists(template_path):
                return None, "Template file not found"
            
            doc = Document(template_path)
            
            # ====== EXTRACT TOP-N CATEGORIES AND SUPPLIERS ======
            top_categories = {}
            top_suppliers = {}
            category_other_total = 0.0
            supplier_other_total = 0.0
            
            if scope31_supplier_specific_results is not None and not scope31_supplier_specific_results.empty:
                # Extract categories
                category_df = scope31_supplier_specific_results[scope31_supplier_specific_results['Supplier'] != 'TOTAL'].copy()
                if 'Spend Category' in category_df.columns:
                    cat_col = 'Spend Category'
                else:
                    cat_col = 'SpendCategory'
                
                category_totals = category_df.groupby(cat_col)['kg_CO2e'].sum().sort_values(ascending=False)
                category_totals_t = category_totals / 1000
                total_cat_emissions = category_totals_t.sum()
                
                top_10_categories = category_totals_t.head(10)
                category_other_total = category_totals_t.iloc[10:].sum() if len(category_totals_t) > 10 else 0.0
                
                for idx, (cat_name, emissions) in enumerate(top_10_categories.items(), 1):
                    pct = (emissions / total_cat_emissions * 100) if total_cat_emissions > 0 else 0
                    top_categories[f'cat1_topcat{idx}_name'] = str(cat_name)
                    top_categories[f'cat1_topcat{idx}_ghg_tco2e'] = f"{emissions:,.2f}"
                    top_categories[f'cat1_topcat{idx}_share_percent'] = f"{pct:.1f}%"
                
                if category_other_total > 0:
                    other_pct = (category_other_total / total_cat_emissions * 100) if total_cat_emissions > 0 else 0
                    top_categories['cat1_topcat_other_ghg_tco2e'] = f"{category_other_total:,.2f}"
                    top_categories['cat1_topcat_other_share_percent'] = f"{other_pct:.1f}%"
                
                # Extract suppliers
                supplier_totals = category_df.groupby('Supplier')['kg_CO2e'].sum().sort_values(ascending=False)
                supplier_totals_t = supplier_totals / 1000
                total_supplier_emissions = supplier_totals_t.sum()
                
                top_11_suppliers_list = supplier_totals_t.head(11)
                supplier_other_total = supplier_totals_t.iloc[11:].sum() if len(supplier_totals_t) > 11 else 0.0
                
                for idx, (supplier_name, emissions) in enumerate(top_11_suppliers_list.items(), 1):
                    pct = (emissions / total_supplier_emissions * 100) if total_supplier_emissions > 0 else 0
                    top_suppliers[f'cat1_topsup{idx}_name'] = str(supplier_name)
                    top_suppliers[f'cat1_topsup{idx}_ghg_tco2e'] = f"{emissions:,.2f}"
                    top_suppliers[f'cat1_topsup{idx}_share_percent'] = f"{pct:.1f}%"
                
                if supplier_other_total > 0:
                    other_pct = (supplier_other_total / total_supplier_emissions * 100) if total_supplier_emissions > 0 else 0
                    top_suppliers['cat1_topsup_other_ghg_tco2e'] = f"{supplier_other_total:,.2f}"
                    top_suppliers['cat1_topsup_other_share_percent'] = f"{other_pct:.1f}%"
            
            # ====== BUILD PLACEHOLDER DICTIONARY ======
            placeholders = {}

            # Main consolidated metrics — keys match template placeholder names exactly
            placeholders['combined_total'] = f"{combined_total:,.2f}"
            placeholders['report_scope31_total'] = f"{report_scope31_total:,.2f}"
            placeholders['report_scope34_total'] = f"{report_scope34_total:,.2f}"
            placeholders['report_spend_amount'] = f"{report_spend_amount:,.2f}" if report_spend_amount is not None else 'N/A'
            placeholders['total_shipment_weight_ton'] = f"{scope34_weight_ton:,.2f}"
            placeholders['total_transport_activity_ton_km'] = f"{scope34_ton_km:,.2f}"
            placeholders['report_po_count'] = str(report_po_count)
            placeholders['cat4_num_shipments'] = str(awb_shipments + bol_shipments)

            # Percentages
            if combined_total > 0:
                placeholders['cat1_total_share_percent'] = f"{(report_scope31_total / combined_total * 100):.1f}%"
                placeholders['cat4_total_share_percent'] = f"{(report_scope34_total / combined_total * 100):.1f}%"
                placeholders['combined_percentage_3'] = f"{(report_scope31_total / combined_total * 100):.1f}%"
                placeholders['combined_percentage_4'] = f"{(report_scope34_total / combined_total * 100):.1f}%"
            else:
                placeholders['cat1_total_share_percent'] = 'N/A'
                placeholders['cat4_total_share_percent'] = 'N/A'
                placeholders['combined_percentage_3'] = 'N/A'
                placeholders['combined_percentage_4'] = 'N/A'

            # Intensities — use the same variables calculated in the transport section
            if report_spend_amount is not None and report_spend_amount > 0:
                _kpi_cat1_intensity = report_scope31_total / report_spend_amount
                placeholders['kpi_cat1_intensity'] = f"{_kpi_cat1_intensity:,.6f}"
            else:
                placeholders['kpi_cat1_intensity'] = 'N/A'

            placeholders['overall_intensity'] = f"{overall_intensity:,.1f}" if overall_intensity is not None else 'N/A'
            placeholders['transport_intensity'] = f"{transport_intensity:,.1f}" if transport_intensity is not None else 'N/A'
            placeholders['air_intensity'] = f"{air_intensity:,.1f}" if air_intensity is not None else 'N/A'
            placeholders['sea_intensity'] = f"{sea_intensity:,.1f}" if sea_intensity is not None else 'N/A'
            placeholders['mixed_mode_tce1_intensity'] = f"{mixed_mode_tce1_intensity:,.1f}" if mixed_mode_tce1_intensity is not None else 'N/A'
            placeholders['hubs_intensity'] = f"{hubs_intensity:,.1f}" if hubs_intensity is not None else 'N/A'

            # Methodology levels (scope 3.1) — keys match template names exactly
            placeholders['scope31_total'] = f"{scope31_total:,.2f}"
            placeholders['scope31_exiobase_total'] = f"{scope31_exiobase_total:,.2f}"
            placeholders['scope31_supplier_total'] = f"{scope31_supplier_total:,.2f}"

            # TCE breakdown
            tce1_total = tce1_results.get('Total', 0.0)
            tce2_total = tce2_results.get('Total', 0.0)
            tce3_total = tce3_value
            tce4_total = tce4_results.get('Total', 0.0)
            scope34_total = scope34_grand_total

            placeholders['tce1_value'] = f"{tce1_total:,.2f}"
            placeholders['tce2_value'] = f"{tce2_total:,.2f}"
            placeholders['tce_3_value'] = f"{tce3_total:,.2f}"
            placeholders['tce4_value'] = f"{tce4_total:,.2f}"

            if scope34_total > 0:
                placeholders['cat4_tce1_share_percent'] = f"{(tce1_total / scope34_total * 100):.1f}%"
                placeholders['cat4_tce2_share_percent'] = f"{(tce2_total / scope34_total * 100):.1f}%"
                placeholders['cat4_tce3_share_percent'] = f"{(tce3_total / scope34_total * 100):.1f}%"
                placeholders['cat4_tce4_share_percent'] = f"{(tce4_total / scope34_total * 100):.1f}%"
            else:
                placeholders['cat4_tce1_share_percent'] = 'N/A'
                placeholders['cat4_tce2_share_percent'] = 'N/A'
                placeholders['cat4_tce3_share_percent'] = 'N/A'
                placeholders['cat4_tce4_share_percent'] = 'N/A'

            # Transport modes — keys match template names exactly
            air_tce3_val = awb_emissions
            sea_tce3_val = bol_emissions
            tce1_modes_val = tce1_results.get('Total', 0.0)
            hub_operations_val = tce2_total + tce4_total

            placeholders['air_tce3'] = f"{air_tce3_val:,.2f}"
            placeholders['sea_tce3'] = f"{sea_tce3_val:,.2f}"
            placeholders['tce1_modes'] = f"{tce1_modes_val:,.2f}"
            placeholders['hub_operations'] = f"{hub_operations_val:,.2f}"

            if scope34_total > 0:
                placeholders['cat4_mode1_share_percent'] = f"{(air_tce3_val / scope34_total * 100):.1f}%"
                placeholders['cat4_mode2_share_percent'] = f"{(sea_tce3_val / scope34_total * 100):.1f}%"
                placeholders['cat4_mode3_share_percent'] = f"{(tce1_modes_val / scope34_total * 100):.1f}%"
                placeholders['cat4_mode4_share_percent'] = f"{(hub_operations_val / scope34_total * 100):.1f}%"
            else:
                placeholders['cat4_mode1_share_percent'] = 'N/A'
                placeholders['cat4_mode2_share_percent'] = 'N/A'
                placeholders['cat4_mode3_share_percent'] = 'N/A'
                placeholders['cat4_mode4_share_percent'] = 'N/A'

            # Top categories and suppliers (merge then rename "other" keys to match template)
            placeholders.update(top_categories)
            placeholders.update(top_suppliers)

            # Rename "other" row keys to match template names
            for old, new in [
                ('cat1_topcat_other_ghg_tco2e', 'cat1_other_categories_ghg_tco2e'),
                ('cat1_topcat_other_share_percent', 'cat1_other_categories_share_percent'),
                ('cat1_topsup_other_ghg_tco2e', 'cat1_other_suppliers_ghg_tco2e'),
                ('cat1_topsup_other_share_percent', 'cat1_other_suppliers_share_percent'),
            ]:
                if old in placeholders:
                    placeholders[new] = placeholders.pop(old)

            # Top category shortcut placeholders
            placeholders['top_category'] = top_categories.get('cat1_topcat1_name', 'N/A')
            placeholders['top_category_pct'] = top_categories.get('cat1_topcat1_share_percent', 'N/A')

            # Add report year
            placeholders['report_year'] = str(report_year)
            
            # ====== REPLACE PLACEHOLDERS IN DOCUMENT ======
            def replace_text_in_paragraph(paragraph, replacements):
                """Replace placeholders preserving per-run formatting (highlights, bold, etc.).

                Modifies run text in-place so that yellow highlights, bold, colour, and
                all other character formatting survive the substitution unchanged.

                For placeholders split across multiple runs (common when Word applies
                spell-check or partial formatting), the pattern is located in the joined
                text, the first involved run receives (prefix + value), the last involved
                run retains any trailing text, and all middle runs are cleared.
                """
                for placeholder, value in replacements.items():
                    pattern = '{{' + placeholder + '}}'
                    runs = paragraph.runs
                    if not runs:
                        continue

                    # Build full text and per-run character ranges
                    full_text = ''
                    run_ranges = []
                    for i, run in enumerate(runs):
                        s = len(full_text)
                        full_text += run.text
                        run_ranges.append((i, s, len(full_text)))

                    if pattern not in full_text:
                        continue

                    pat_start = full_text.find(pattern)
                    pat_end = pat_start + len(pattern)

                    # Identify every run that overlaps with the pattern
                    involved = [(i, s, e) for i, s, e in run_ranges if s < pat_end and e > pat_start]
                    if not involved:
                        continue

                    first_i, first_s, _ = involved[0]
                    last_i, _, last_e = involved[-1]

                    # Text in the first run before the pattern starts
                    prefix = full_text[first_s:pat_start]
                    # Text in the last run after the pattern ends
                    suffix = full_text[pat_end:last_e]

                    if first_i == last_i:
                        # Entire pattern in one run — replace text, keep all formatting
                        runs[first_i].text = prefix + value + suffix
                    else:
                        # Pattern spans multiple runs — first run gets prefix+value,
                        # last run retains its suffix (and its own formatting),
                        # middle runs are cleared
                        runs[first_i].text = prefix + value
                        runs[last_i].text = suffix
                        for i, _, _ in involved[1:-1]:
                            runs[i].text = ''
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, placeholders)
            
            # ====== GENERATE OUTPUT ======
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output, None
        
        except Exception as e:
            return None, f"Error generating report: {str(e)}"
    
    # Add download button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("📥 Download Populated Report (Word)", use_container_width=True, key="download_report"):
            doc_bytes, error = generate_word_report()
            if error:
                st.error(error)
            elif doc_bytes:
                st.download_button(
                    label="📄 Download PAHO GHG Report",
                    data=doc_bytes,
                    file_name=f"PAHO_GHG_Report_{report_year}_Scope3_Populated.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_button_doc"
                )
                st.success("✅ Report generated successfully! Click the button above to download.")
            else:
                st.error("Failed to generate report")


if __name__ == "__main__":
    main()
